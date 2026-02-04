#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-fill-template
# @raycast.mode fullOutput
# @raycast.icon 📋
# @raycast.packageName DOCX
# @raycast.description Word 格式刷工具（从模板提取格式应用到内容）
# -*- coding: utf-8 -*-
"""
docx_fill_template.py - Word 格式刷工具

功能：从模板提取格式（标题1、标题2、正文等样式），应用到内容上生成新文档

用法：
    python docx_fill_template.py 内容.md 模板.docx
    python docx_fill_template.py 内容.md 模板.docx -o 输出.docx

格式映射：
    # xxx      → Title（封面标题）
    ## xxx     → Heading 1（标题1）
    ### xxx    → Heading 2（标题2）
    #### xxx   → Heading 3（标题3）
    普通文本    → Normal（正文）

示例：
    python docx_fill_template.py 说明书内容.md 格式模板.docx -o 说明书.docx
"""

import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt
from copy import deepcopy


def parse_content(content_path: Path) -> list:
    """
    解析内容文件，返回 [(样式名, 文本), ...]
    
    规则：
        # xxx    → Title
        ## xxx   → Heading 1
        ### xxx  → Heading 2
        #### xxx → Heading 3
        空行      → 段落分隔
        普通文本  → Normal
    """
    with open(content_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')
    
    paragraphs = []
    current_body = []
    
    for line in lines:
        # 检测标题
        if line.startswith('#### '):
            # 先保存累积的正文
            if current_body:
                paragraphs.append(('Normal', ' '.join(current_body)))
                current_body = []
            text = line[5:].strip()
            if text:
                paragraphs.append(('Heading 3', text))
        elif line.startswith('### '):
            if current_body:
                paragraphs.append(('Normal', ' '.join(current_body)))
                current_body = []
            text = line[4:].strip()
            if text:
                paragraphs.append(('Heading 2', text))
        elif line.startswith('## '):
            if current_body:
                paragraphs.append(('Normal', ' '.join(current_body)))
                current_body = []
            text = line[3:].strip()
            if text:
                paragraphs.append(('Heading 1', text))
        elif line.startswith('# '):
            if current_body:
                paragraphs.append(('Normal', ' '.join(current_body)))
                current_body = []
            text = line[2:].strip()
            if text:
                paragraphs.append(('Title', text))
        elif not line.strip():
            # 空行：保存累积的正文
            if current_body:
                paragraphs.append(('Normal', ' '.join(current_body)))
                current_body = []
        else:
            # 正文行
            current_body.append(line.strip())
    
    # 保存最后的正文
    if current_body:
        paragraphs.append(('Normal', ' '.join(current_body)))
    
    return paragraphs


def copy_style_properties(source_style, target_para):
    """
    复制样式属性到段落
    """
    # 应用样式名
    try:
        target_para.style = source_style
    except:
        pass


def generate_document(template_path: Path, content: list, output_path: Path, verbose: bool = True):
    """
    用模板格式生成新文档
    
    Args:
        template_path: 模板文件路径
        content: [(样式名, 文本), ...] 内容列表
        output_path: 输出文件路径
        verbose: 是否打印详细信息
    """
    # 打开模板
    template_doc = Document(template_path)
    
    # 创建新文档（基于模板，继承样式）
    # 方法：复制模板，清空内容，添加新内容
    doc = Document(template_path)
    
    # 清空所有段落内容（保留样式定义）
    # 注意：不能直接删除段落，会破坏样式引用
    # 我们创建一个全新的文档，但复制模板的样式
    
    # 实际上更好的方法是：创建空文档，复制样式
    new_doc = Document()
    
    # 复制模板的样式到新文档
    # python-docx 的样式是基于 XML 的，直接复制比较复杂
    # 更简单的方法：使用模板作为基础，清空内容
    
    doc = Document(template_path)
    
    # 删除所有现有段落（除了必要的）
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)
    
    # 删除所有表格
    for table in list(doc.tables):
        t = table._element
        t.getparent().remove(t)
    
    # 统计
    style_counts = {}
    
    # 添加新内容
    for style_name, text in content:
        try:
            para = doc.add_paragraph(text, style=style_name)
            style_counts[style_name] = style_counts.get(style_name, 0) + 1
        except KeyError:
            # 如果样式不存在，使用 Normal
            if verbose:
                print(f"  ⚠️ 样式 '{style_name}' 不存在，使用 Normal")
            para = doc.add_paragraph(text, style='Normal')
            style_counts['Normal'] = style_counts.get('Normal', 0) + 1
    
    # 保存
    doc.save(output_path)
    
    if verbose:
        print(f"\n📊 样式统计:")
        for style, count in sorted(style_counts.items()):
            print(f"   {style}: {count} 个")
        print(f"\n✅ 完成!")
        print(f"📁 输出: {output_path}")
    
    return style_counts


def main():
    parser = argparse.ArgumentParser(
        description='Word 格式刷工具 - 从模板提取格式，应用到内容',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
格式映射:
  # xxx      → Title（封面标题）
  ## xxx     → Heading 1（标题1）
  ### xxx    → Heading 2（标题2）
  #### xxx   → Heading 3（标题3）
  普通文本    → Normal（正文）

示例:
  %(prog)s 说明书内容.md 格式模板.docx
  %(prog)s 说明书内容.md 格式模板.docx -o 说明书.docx
        """
    )
    
    parser.add_argument('content', type=Path, help='内容文件路径 (.md/.txt)')
    parser.add_argument('template', type=Path, help='模板文件路径 (.docx)')
    parser.add_argument('-o', '--output', type=Path, default=None, 
                        help='输出文件路径（默认: 内容文件名_格式化.docx）')
    parser.add_argument('-q', '--quiet', action='store_true', help='静默模式')
    
    args = parser.parse_args()
    
    # 检查文件
    if not args.content.exists():
        print(f"❌ 错误: 内容文件不存在: {args.content}")
        sys.exit(1)
    
    if not args.template.exists():
        print(f"❌ 错误: 模板文件不存在: {args.template}")
        sys.exit(1)
    
    # 输出路径
    if args.output:
        output_path = args.output
    else:
        output_path = args.content.parent / f"{args.content.stem}_格式化.docx"
    
    verbose = not args.quiet
    
    if verbose:
        print(f"📄 内容文件: {args.content}")
        print(f"📋 格式模板: {args.template}")
    
    # 解析内容
    content = parse_content(args.content)
    
    if verbose:
        print(f"📝 解析到 {len(content)} 个段落")
    
    # 生成文档
    generate_document(args.template, content, output_path, verbose)


if __name__ == "__main__":
    main()
