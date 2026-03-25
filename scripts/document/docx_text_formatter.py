#!/usr/bin/env python3
"""
文本格式自动修复工具 - DOCX版本
支持处理Word文档中的格式问题

功能列表：
1. 双引号格式修复
   将所有双引号统一为中文标准引号："和"
   奇数个引号 → " (中文左双引号 U+201C)
   偶数个引号 → " (中文右双引号 U+201D)

2. 英文标点符号转中文
   , → ，  (逗号)
   : → ：  (冒号)
   ; → ；  (分号)
   ! → ！  (感叹号)
   ? → ？  (问号)
   ( → （  (左括号)
   ) → ）  (右括号)

3. 中文单位转标准符号
   面积：平方米 → m²、平方公里 → km²
   体积：立方米 → m³、立方厘米 → cm³
   长度：公里 → km、厘米 → cm、毫米 → mm
   质量：公斤 → kg、毫克 → mg
   容量：毫升 → mL、微升 → μL
   时间：小时 → h、分钟 → min、秒钟 → s
   温度：摄氏度 → ℃、华氏度 → ℉

4. 单位上标格式化
   m2 → m²、m3 → m³、km2 → km²、km3 → km³

使用方法：
    python3 text_formatter_docx.py 文件名.docx
"""

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent.parent.parent / “lib”))
from finder import get_input_files
from progress import ProgressTracker
from text_format_rules import fix_punctuation, fix_quotes, fix_units

# ===== 配置选项 =====
SKIP_FOOTER = True


QUOTE_CHARS = {"\u201c", "\u201d"}
QUOTE_FONT = "\u5b8b\u4f53"  # 宋体


def _set_run_font_songti(run_element):
    """为 run 的 rPr 设置宋体字体（ascii + hAnsi + eastAsia）"""
    rPr = run_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), QUOTE_FONT)
    rFonts.set(qn("w:hAnsi"), QUOTE_FONT)
    rFonts.set(qn("w:eastAsia"), QUOTE_FONT)
    rFonts.set(qn("w:hint"), "eastAsia")


def _split_run_at_quotes(run):
    """
    将 run 在引号位置拆分，返回 [(text, is_quote), ...] 片段列表。
    如果没有引号，返回 None。
    """
    text = run.text
    if not text or not any(c in QUOTE_CHARS for c in text):
        return None

    segments = []
    buf = []
    for c in text:
        if c in QUOTE_CHARS:
            if buf:
                segments.append(("".join(buf), False))
                buf = []
            segments.append((c, True))
        else:
            buf.append(c)
    if buf:
        segments.append(("".join(buf), False))
    return segments


def _apply_quote_split(run, segments):
    """
    根据 segments 拆分 run：第一段复用原 run，后续段插入新 run。
    引号段设置宋体。
    """
    parent = run._element.getparent()

    # 第一段复用原 run
    first_text, first_is_quote = segments[0]
    run.text = first_text
    if first_is_quote:
        _set_run_font_songti(run._element)

    # 后续段：复制原 run 的格式，插入到原 run 之后
    insert_after = run._element
    for seg_text, is_quote in segments[1:]:
        new_r = copy.deepcopy(run._element)
        # 设置文本（清除 deepcopy 带来的旧文本）
        for t_elem in new_r.findall(qn("w:t")):
            new_r.remove(t_elem)
        t_elem = OxmlElement("w:t")
        t_elem.text = seg_text
        # 保留空格
        t_elem.set(qn("xml:space"), "preserve")
        new_r.append(t_elem)

        if is_quote:
            _set_run_font_songti(new_r)
        else:
            # 非引号段恢复原 run 的字体（去掉宋体覆盖）
            rPr = new_r.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                orig_rPr = run._element.find(qn("w:rPr"))
                orig_rFonts = orig_rPr.find(qn("w:rFonts")) if orig_rPr is not None else None
                if rFonts is not None and orig_rFonts is not None:
                    # 用原始字体信息覆盖
                    rPr.replace(rFonts, copy.deepcopy(orig_rFonts))
                elif rFonts is not None and orig_rFonts is None:
                    rPr.remove(rFonts)

        parent.insert(list(parent).index(insert_after) + 1, new_r)
        insert_after = new_r


def process_paragraph(paragraph, stats):
    """
    处理段落中的文本，逐run处理以保持每个run的字体格式。
    引号拆分为独立run并设置宋体。
    """
    if not paragraph.runs:
        return

    # 引号计数器跨run维护，保证配对正确
    quote_counter = 0

    # 先收集需要处理的 runs（遍历中会插入新 run，所以先快照）
    original_runs = list(paragraph.runs)

    for run in original_runs:
        if not run.text:
            continue

        original_text = run.text

        # 应用所有转换
        fixed_text, quote_count, quote_counter = fix_quotes(original_text, quote_counter)
        fixed_text, punct_count = fix_punctuation(fixed_text)
        fixed_text, unit_count = fix_units(fixed_text)

        # 更新统计
        stats["quotes"] += quote_count
        stats["punctuation"] += punct_count
        stats["units"] += unit_count

        if fixed_text != original_text:
            run.text = fixed_text

        # 拆分引号为独立 run 并设置宋体
        segments = _split_run_at_quotes(run)
        if segments:
            _apply_quote_split(run, segments)


def process_table(table, stats):
    """
    处理表格中的文本
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph, stats)


def process_docx(input_file):
    """
    处理DOCX文件
    """
    input_path = Path(input_file)

    if not input_path.exists():
        print(f"❌ 错误：文件不存在 - {input_file}")
        return False

    if input_path.suffix.lower() != ".docx":
        print("❌ 错误：文件必须是.docx格式")
        return False

    # 生成输出文件名
    output_path = input_path.parent / f"{input_path.stem}_fixed{input_path.suffix}"

    try:
        # 读取文档
        print(f"📖 正在读取文件: {input_path.name}")
        doc = Document(input_path)

        # 统计信息
        stats = {"quotes": 0, "punctuation": 0, "units": 0}

        # 处理所有段落
        print("🔄 正在处理段落...")
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph, stats)

        # 处理所有表格
        print("🔄 正在处理表格...")
        for table in doc.tables:
            process_table(table, stats)

        # 处理页眉页脚
        if SKIP_FOOTER:
            # 彻底删除页眉页脚
            print("🗑️  正在删除页眉页脚...")
            from docx.oxml.ns import qn

            for section in doc.sections:
                sectPr = section._sectPr
                # 删除所有页眉引用
                for headerRef in sectPr.findall(qn("w:headerReference")):
                    sectPr.remove(headerRef)
                # 删除所有页脚引用
                for footerRef in sectPr.findall(qn("w:footerReference")):
                    sectPr.remove(footerRef)
        else:
            # 处理页眉页脚（格式修复）
            print("🔄 正在处理页眉页脚...")
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        process_paragraph(paragraph, stats)
                    for table in section.header.tables:
                        process_table(table, stats)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        process_paragraph(paragraph, stats)
                    for table in section.footer.tables:
                        process_table(table, stats)

        # 保存文件
        print("💾 正在保存文件...")
        doc.save(output_path)

        print("✅ 处理完成！")
        print(f"   - 共替换了 {stats['quotes']} 个引号")
        print(f"   - 共替换了 {stats['punctuation']} 个标点符号")
        print(f"   - 共转换了 {stats['units']} 个单位")
        print(f"   - 输出文件: {output_path.name}")

        return True

    except Exception as e:
        print(f"❌ 处理失败：{e}")
        import traceback

        traceback.print_exc()
        return False


if __name__ == "__main__":
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext="docx")

    if not files:
        print("❌ 错误：缺少文件名参数")
        print("\n使用方法：")
        print("  1. 在 Finder 中选中 .docx 文件，然后运行此脚本")
        print("  2. 或在命令行中提供文件路径")
        print("\n示例：")
        print("    python3 docx_text_formatter.py 文件名.docx")
        print("    python3 docx_text_formatter.py file1.docx file2.docx")
        sys.exit(1)

    print("=" * 50)
    print("文本格式自动修复工具 - DOCX版本")
    print("=" * 50)

    tracker = ProgressTracker()

    for file_path in files:
        print(f"\n处理文件: {Path(file_path).name}")
        success = process_docx(str(file_path))
        if success:
            tracker.add_success()
        else:
            tracker.add_error()

    print("\n" + "=" * 50)
    tracker.show_summary("文件处理")
