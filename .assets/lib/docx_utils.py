#!/usr/bin/env python3
"""
Word 文档操作公共模块
提供 docx XML 操作的公共函数
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from typing import Optional


# ===== XML 元素操作 =====

def create_element(name: str) -> OxmlElement:
    """创建 XML 元素"""
    return OxmlElement(name)


def create_attribute(element: OxmlElement, name: str, value: str):
    """设置 XML 属性"""
    element.set(qn(name), value)


# ===== 字体设置 =====

def set_run_font(run, font_name: str = '宋体', font_size: float = None, 
                 bold: bool = None, east_asia: str = None):
    """
    设置 run 的字体属性
    
    Args:
        run: docx Run 对象
        font_name: 西文字体名称
        font_size: 字号（磅）
        bold: 是否加粗
        east_asia: 东亚字体名称（中文）
    """
    if font_name:
        run.font.name = font_name
    if east_asia:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    elif font_name:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.font.bold = bold


# ===== 页码域 =====

def add_page_number(paragraph):
    """
    在段落中添加页码域
    
    Args:
        paragraph: Paragraph 对象
    """
    run = paragraph.add_run()
    
    # fldChar begin
    fldChar_begin = create_element('w:fldChar')
    create_attribute(fldChar_begin, 'w:fldCharType', 'begin')
    
    # instrText
    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    
    # fldChar separate
    fldChar_separate = create_element('w:fldChar')
    create_attribute(fldChar_separate, 'w:fldCharType', 'separate')
    
    # 页码文本占位符
    page_num_run = create_element('w:t')
    page_num_run.text = "1"
    
    # fldChar end
    fldChar_end = create_element('w:fldChar')
    create_attribute(fldChar_end, 'w:fldCharType', 'end')
    
    # 组装域代码
    run._r.append(fldChar_begin)
    run2 = paragraph.add_run()
    run2._r.append(instrText)
    run3 = paragraph.add_run()
    run3._r.append(fldChar_separate)
    run4 = paragraph.add_run()
    run4._r.append(page_num_run)
    run5 = paragraph.add_run()
    run5._r.append(fldChar_end)


# ===== 样式引用域 =====

def add_styleref_field(paragraph, style_name: str = "标题 1"):
    """
    在段落中添加 STYLEREF 域（引用指定样式的文本）
    
    Args:
        paragraph: Paragraph 对象
        style_name: 样式名称
    """
    run = paragraph.add_run()
    
    fldChar_begin = create_element('w:fldChar')
    create_attribute(fldChar_begin, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' STYLEREF "{style_name}" '
    
    fldChar_separate = create_element('w:fldChar')
    create_attribute(fldChar_separate, 'w:fldCharType', 'separate')
    
    ref_text = create_element('w:t')
    ref_text.text = ""
    
    fldChar_end = create_element('w:fldChar')
    create_attribute(fldChar_end, 'w:fldCharType', 'end')
    
    run._r.append(fldChar_begin)
    run2 = paragraph.add_run()
    run2._r.append(instrText)
    run3 = paragraph.add_run()
    run3._r.append(fldChar_separate)
    run4 = paragraph.add_run()
    run4._r.append(ref_text)
    run5 = paragraph.add_run()
    run5._r.append(fldChar_end)


# ===== 分节符 =====

def insert_section_break(doc, para_index: int, break_type: str = 'nextPage') -> bool:
    """
    在指定段落前插入分节符
    
    Args:
        doc: Document 对象
        para_index: 段落索引
        break_type: 分节类型 ('nextPage', 'continuous', 'evenPage', 'oddPage')
    
    Returns:
        是否成功
    """
    if para_index <= 0 or para_index >= len(doc.paragraphs):
        return False
    
    prev_para = doc.paragraphs[para_index - 1]
    pPr = prev_para._element.get_or_add_pPr()
    
    sectPr = create_element('w:sectPr')
    sectType = create_element('w:type')
    create_attribute(sectType, 'w:val', break_type)
    sectPr.append(sectType)
    
    # 复制页面设置
    if doc.sections:
        current_section = doc.sections[0]
        
        pgSz = create_element('w:pgSz')
        create_attribute(pgSz, 'w:w', str(int(current_section.page_width.twips)))
        create_attribute(pgSz, 'w:h', str(int(current_section.page_height.twips)))
        sectPr.append(pgSz)
        
        pgMar = create_element('w:pgMar')
        create_attribute(pgMar, 'w:top', str(int(current_section.top_margin.twips)))
        create_attribute(pgMar, 'w:right', str(int(current_section.right_margin.twips)))
        create_attribute(pgMar, 'w:bottom', str(int(current_section.bottom_margin.twips)))
        create_attribute(pgMar, 'w:left', str(int(current_section.left_margin.twips)))
        create_attribute(pgMar, 'w:header', '720')
        create_attribute(pgMar, 'w:footer', '720')
        sectPr.append(pgMar)
    
    pPr.append(sectPr)
    return True


# ===== 查找段落 =====

def find_paragraph_by_text(doc, patterns: list[str]) -> int:
    """
    查找包含指定文本的段落索引
    
    Args:
        doc: Document 对象
        patterns: 文本模式列表
    
    Returns:
        段落索引，未找到返回 -1
    """
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for pattern in patterns:
            if text == pattern or text.startswith(pattern):
                return i
    return -1


# ===== 获取第一个标题 =====

def get_first_heading_text(doc, style_prefix: str = "Heading") -> Optional[str]:
    """
    获取第一个标题段落的文本
    
    Args:
        doc: Document 对象
        style_prefix: 标题样式前缀
    
    Returns:
        标题文本，未找到返回 None
    """
    for para in doc.paragraphs:
        if para.style and para.style.name:
            if para.style.name.startswith(style_prefix) or para.style.name.startswith("标题"):
                text = para.text.strip()
                if text:
                    return text
    return None

