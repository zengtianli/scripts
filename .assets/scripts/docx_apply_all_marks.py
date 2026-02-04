#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-all-marks
# @raycast.mode fullOutput
# @raycast.icon 📄
# @raycast.packageName Scripts
# @raycast.description Apply header footer watermark
"""
页眉页脚水印设置工具
为Word文档添加标准的页眉、页脚和水印

功能：
1. 页脚：公司名称（左对齐）+ 页码（居中）
2. 页眉：文档标题（左对齐）+ 章节引用（右对齐，自动跟随标题1）
3. 水印：ZDWP 对角线水印，Arial 96号

用法:
    python3 apply_header_footer_watermark.py <input.docx> [文档标题]
    
示例:
    python3 apply_header_footer_watermark.py document.docx
    python3 apply_header_footer_watermark.py document.docx "浙江省水利工程报告"
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import copy

from common_utils import get_input_files

# 公司名称
COMPANY_NAME = "浙江省水利水电勘测设计院有限责任公司"

# 水印设置
WATERMARK_TEXT = "ZDWP"
WATERMARK_FONT = "Arial"
WATERMARK_SIZE = 96


def get_first_heading1_text(doc):
    """
    获取文档中第一个标题1的文本
    
    Args:
        doc: Document对象
        
    Returns:
        str: 标题文本，如果没有找到返回空字符串
    """
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name
            if 'Heading 1' in style_name or '标题 1' in style_name or style_name == '标题1':
                text = paragraph.text.strip()
                if text:
                    return text
    return ""


def create_element(name):
    """创建XML元素"""
    return OxmlElement(name)


def create_attribute(element, name, value):
    """设置XML属性"""
    element.set(qn(name), value)


def add_page_number(paragraph):
    """
    在段落中添加页码域
    
    Args:
        paragraph: Paragraph对象
    """
    run = paragraph.add_run()
    
    # 创建 fldChar begin
    fldChar_begin = create_element('w:fldChar')
    create_attribute(fldChar_begin, 'w:fldCharType', 'begin')
    
    # 创建 instrText (PAGE)
    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    
    # 创建 fldChar separate
    fldChar_separate = create_element('w:fldChar')
    create_attribute(fldChar_separate, 'w:fldCharType', 'separate')
    
    # 创建页码文本占位符
    page_num_run = create_element('w:t')
    page_num_run.text = "1"
    
    # 创建 fldChar end
    fldChar_end = create_element('w:fldChar')
    create_attribute(fldChar_end, 'w:fldCharType', 'end')
    
    # 组装域代码
    run._r.append(fldChar_begin)
    
    run2 = paragraph.add_run()
    run2._r.append(instrText)
    
    run3 = paragraph.add_run()
    run3._r.append(fldChar_separate)
    
    run4 = paragraph.add_run()
    run4._r.append(page_num_run)
    
    run5 = paragraph.add_run()
    run5._r.append(fldChar_end)


def add_styleref_field(paragraph, style_name="标题 1"):
    """
    在段落中添加STYLEREF域（引用指定样式的文本）
    
    Args:
        paragraph: Paragraph对象
        style_name: 样式名称（默认"标题 1"）
    """
    run = paragraph.add_run()
    
    # 创建 fldChar begin
    fldChar_begin = create_element('w:fldChar')
    create_attribute(fldChar_begin, 'w:fldCharType', 'begin')
    
    # 创建 instrText (STYLEREF)
    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' STYLEREF "{style_name}" \\* MERGEFORMAT '
    
    # 创建 fldChar separate
    fldChar_separate = create_element('w:fldChar')
    create_attribute(fldChar_separate, 'w:fldCharType', 'separate')
    
    # 创建占位符文本
    placeholder = create_element('w:t')
    placeholder.text = "章节标题"
    
    # 创建 fldChar end
    fldChar_end = create_element('w:fldChar')
    create_attribute(fldChar_end, 'w:fldCharType', 'end')
    
    # 组装域代码
    run._r.append(fldChar_begin)
    
    run2 = paragraph.add_run()
    run2._r.append(instrText)
    
    run3 = paragraph.add_run()
    run3._r.append(fldChar_separate)
    
    run4 = paragraph.add_run()
    run4._r.append(placeholder)
    
    run5 = paragraph.add_run()
    run5._r.append(fldChar_end)


def set_footer(doc, company_name=COMPANY_NAME):
    """
    设置页脚：公司名称（左）+ 页码（居中）
    
    Args:
        doc: Document对象
        company_name: 公司名称
    """
    # 确保文档有节
    section = doc.sections[0]
    
    # 获取或创建页脚
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # 清空现有内容
    for paragraph in footer.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    # 创建页脚段落
    para = footer.add_paragraph()
    
    # 使用制表位来实现左对齐和居中
    # 设置段落制表位
    pPr = para._element.get_or_add_pPr()
    tabs = create_element('w:tabs')
    
    # 居中制表位（页面中央）
    tab_center = create_element('w:tab')
    create_attribute(tab_center, 'w:val', 'center')
    create_attribute(tab_center, 'w:pos', '4680')  # 约8.25cm，A4纸中央
    tabs.append(tab_center)
    
    # 右对齐制表位
    tab_right = create_element('w:tab')
    create_attribute(tab_right, 'w:val', 'right')
    create_attribute(tab_right, 'w:pos', '9360')  # 约16.5cm
    tabs.append(tab_right)
    
    pPr.append(tabs)
    
    # 添加公司名称
    run_company = para.add_run(company_name)
    run_company.font.size = Pt(9)
    run_company.font.name = '宋体'
    run_company._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加制表符
    para.add_run('\t')
    
    # 添加"第"
    run_pre = para.add_run('第 ')
    run_pre.font.size = Pt(9)
    run_pre.font.name = 'Times New Roman'
    
    # 添加页码域
    add_page_number(para)
    
    # 添加"页"
    run_post = para.add_run(' 页')
    run_post.font.size = Pt(9)
    run_post.font.name = '宋体'
    run_post._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_header(doc, doc_title=""):
    """
    设置页眉：文档标题（左）+ 章节引用（右）
    
    Args:
        doc: Document对象
        doc_title: 文档标题，如果为空则自动获取第一个标题1
    """
    # 如果没有提供标题，尝试获取第一个标题1
    if not doc_title:
        doc_title = get_first_heading1_text(doc)
    
    if not doc_title:
        doc_title = "文档标题"
    
    # 确保文档有节
    section = doc.sections[0]
    
    # 获取或创建页眉
    header = section.header
    header.is_linked_to_previous = False
    
    # 清空现有内容
    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    # 创建页眉段落
    para = header.add_paragraph()
    
    # 设置段落制表位
    pPr = para._element.get_or_add_pPr()
    tabs = create_element('w:tabs')
    
    # 右对齐制表位
    tab_right = create_element('w:tab')
    create_attribute(tab_right, 'w:val', 'right')
    create_attribute(tab_right, 'w:pos', '9360')  # 约16.5cm
    tabs.append(tab_right)
    
    pPr.append(tabs)
    
    # 添加文档标题（左对齐）
    run_title = para.add_run(doc_title)
    run_title.font.size = Pt(9)
    run_title.font.name = '宋体'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加制表符
    para.add_run('\t')
    
    # 添加章节引用域（右对齐）
    # 尝试使用不同的标题样式名称
    add_styleref_field(para, "标题 1")
    
    # 添加页眉下划线
    pBdr = create_element('w:pBdr')
    bottom = create_element('w:bottom')
    create_attribute(bottom, 'w:val', 'single')
    create_attribute(bottom, 'w:sz', '4')
    create_attribute(bottom, 'w:space', '1')
    create_attribute(bottom, 'w:color', '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_watermark(doc, text=WATERMARK_TEXT, font=WATERMARK_FONT, size=WATERMARK_SIZE):
    """
    添加对角线文字水印
    使用 lxml 直接创建 VML 元素
    
    Args:
        doc: Document对象
        text: 水印文字
        font: 字体名称
        size: 字号
    """
    from lxml import etree
    
    # VML 命名空间
    NSMAP = {
        'v': 'urn:schemas-microsoft-com:vml',
        'o': 'urn:schemas-microsoft-com:office:office',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }
    
    # 为所有节添加水印
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        
        # 确保有段落
        if not header.paragraphs:
            header.add_paragraph()
        
        # 获取 header 的 XML
        hdr = header._element
        
        # 查找第一个段落
        ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        p_elements = hdr.findall(f'.//{ns_w}p')
        if not p_elements:
            continue
        
        p = p_elements[0]
        
        # 创建 r 元素
        r = etree.SubElement(p, f'{ns_w}r')
        
        # 创建 pict 元素
        pict = etree.SubElement(r, f'{ns_w}pict')
        
        # 创建 VML shape 元素
        ns_v = '{urn:schemas-microsoft-com:vml}'
        ns_o = '{urn:schemas-microsoft-com:office:office}'
        
        shape = etree.SubElement(pict, f'{ns_v}shape')
        shape.set('id', 'PowerPlusWaterMarkObject')
        shape.set(f'{ns_o}spid', '_x0000_s2049')
        shape.set('type', '#_x0000_t136')
        shape.set('style', 
                  'position:absolute;'
                  'margin-left:0;'
                  'margin-top:0;'
                  'width:500pt;'
                  'height:180pt;'
                  'z-index:-251656192;'
                  'mso-wrap-edited:f;'
                  'mso-position-horizontal:center;'
                  'mso-position-horizontal-relative:margin;'
                  'mso-position-vertical:center;'
                  'mso-position-vertical-relative:margin;'
                  'rotation:-45')
        shape.set(f'{ns_o}allowincell', 'f')
        shape.set('fillcolor', 'silver')
        shape.set('stroked', 'f')
        
        # 添加 fill 元素（设置透明度）
        fill = etree.SubElement(shape, f'{ns_v}fill')
        fill.set('opacity', '.3')
        
        # 添加 textpath 元素（设置文字）
        textpath = etree.SubElement(shape, f'{ns_v}textpath')
        textpath.set('style', f'font-family:"{font}";font-size:{size}pt')
        textpath.set('string', text)


def apply_header_footer_watermark(input_file, doc_title=""):
    """
    应用页眉、页脚和水印
    
    Args:
        input_file: 输入文件路径
        doc_title: 文档标题（可选，默认自动获取）
    """
    input_path = Path(input_file)
    
    # 检查文件是否存在
    if not input_path.exists():
        print(f"❌ 错误: 文件不存在: {input_file}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"❌ 错误: 只支持 .docx 文件")
        sys.exit(1)
    
    print(f"🔄 正在处理文件: {input_path.name}")
    
    # 加载文档
    try:
        doc = Document(str(input_path))
    except Exception as e:
        print(f"❌ 错误: 无法打开文件: {e}")
        sys.exit(1)
    
    # 备份原文件
    backup_path = input_path.with_suffix('.docx.backup')
    try:
        import shutil
        shutil.copy2(str(input_path), str(backup_path))
        print(f"ℹ️ 已备份原文件: {backup_path.name}")
    except Exception as e:
        print(f"⚠️ 备份失败: {e}")
    
    # 1. 设置页脚
    print(f"🔄 正在设置页脚...")
    try:
        set_footer(doc, COMPANY_NAME)
        print(f"✅ 页脚设置完成: {COMPANY_NAME} + 页码")
    except Exception as e:
        print(f"⚠️ 页脚设置失败: {e}")
    
    # 2. 设置页眉
    print(f"🔄 正在设置页眉...")
    try:
        if not doc_title:
            doc_title = get_first_heading1_text(doc)
            if doc_title:
                print(f"ℹ️ 自动获取文档标题: {doc_title[:30]}...")
            else:
                doc_title = "文档标题"
                print(f"ℹ️ 未找到标题1，使用默认标题")
        
        set_header(doc, doc_title)
        print(f"✅ 页眉设置完成: 标题 + 章节引用")
    except Exception as e:
        print(f"⚠️ 页眉设置失败: {e}")
    
    # 3. 添加水印
    print(f"🔄 正在添加水印...")
    try:
        add_watermark(doc, WATERMARK_TEXT, WATERMARK_FONT, WATERMARK_SIZE)
        print(f"✅ 水印添加完成: {WATERMARK_TEXT} ({WATERMARK_FONT} {WATERMARK_SIZE}pt, 对角线)")
    except Exception as e:
        print(f"⚠️ 水印添加失败: {e}")
    
    # 保存文档
    try:
        doc.save(str(input_path))
        print(f"\n✅ 处理完成!")
        print(f"   - 已保存: {input_path.name}")
        if backup_path.exists():
            print(f"ℹ️ 如需恢复，请使用备份文件: {backup_path.name}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")
        sys.exit(1)


def main():
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext='docx', allow_multiple=False)
    
    if not files:
        print("页眉页脚水印设置工具")
        print("\n用法: python3 apply_header_footer_watermark.py <input.docx> [文档标题]")
        print("      或在 Finder 中选择 .docx 文件后运行")
        print("\n示例:")
        print("  python3 apply_header_footer_watermark.py document.docx")
        print("  python3 apply_header_footer_watermark.py document.docx \"浙江省水利工程报告\"")
        print("\n功能:")
        print("  1. 页脚: 浙江省水利水电勘测设计院有限责任公司 + 页码")
        print("  2. 页眉: 文档标题（左）+ 章节引用（右，自动跟随标题1）")
        print("  3. 水印: ZDWP 对角线，Arial 96号")
        sys.exit(1)
    
    input_file = files[0]
    doc_title = sys.argv[2] if len(sys.argv) > 2 else ""
    
    apply_header_footer_watermark(input_file, doc_title)


if __name__ == "__main__":
    main()

