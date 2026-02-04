#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-footer
# @raycast.mode fullOutput
# @raycast.icon 📄
# @raycast.packageName Scripts
# @raycast.description Apply footer to Word
"""
页脚设置工具
为Word文档添加标准页脚：公司名称（左对齐）+ 页码（居中）
支持从"前言"开始编号

用法:
    python3 apply_footer.py <input.docx> [公司名称]
    
示例:
    python3 apply_footer.py document.docx
    python3 apply_footer.py document.docx "浙江省水利水电勘测设计院有限责任公司"
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

from common_utils import get_input_files

# 默认公司名称
DEFAULT_COMPANY_NAME = "浙江省水利水电勘测设计院有限责任公司"

# 前言标题匹配模式（支持"前言"、"前  言"等）
PREFACE_PATTERNS = ['前言', '前  言', '前 言']


def create_element(name):
    """创建XML元素"""
    return OxmlElement(name)


def create_attribute(element, name, value):
    """设置XML属性"""
    element.set(qn(name), value)


def add_page_number(paragraph):
    """
    在段落中添加页码域
    
    Args:
        paragraph: Paragraph对象
    """
    run = paragraph.add_run()
    
    # 创建 fldChar begin
    fldChar_begin = create_element('w:fldChar')
    create_attribute(fldChar_begin, 'w:fldCharType', 'begin')
    
    # 创建 instrText (PAGE)
    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    
    # 创建 fldChar separate
    fldChar_separate = create_element('w:fldChar')
    create_attribute(fldChar_separate, 'w:fldCharType', 'separate')
    
    # 创建页码文本占位符
    page_num_run = create_element('w:t')
    page_num_run.text = "1"
    
    # 创建 fldChar end
    fldChar_end = create_element('w:fldChar')
    create_attribute(fldChar_end, 'w:fldCharType', 'end')
    
    # 组装域代码
    run._r.append(fldChar_begin)
    
    run2 = paragraph.add_run()
    run2._r.append(instrText)
    
    run3 = paragraph.add_run()
    run3._r.append(fldChar_separate)
    
    run4 = paragraph.add_run()
    run4._r.append(page_num_run)
    
    run5 = paragraph.add_run()
    run5._r.append(fldChar_end)


def find_preface_paragraph(doc):
    """
    查找"前言"段落的索引
    
    Args:
        doc: Document对象
        
    Returns:
        int: 前言段落的索引，未找到返回-1
    """
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        # 检查是否匹配前言模式
        for pattern in PREFACE_PATTERNS:
            if text == pattern or text.startswith(pattern):
                return i
    return -1


def insert_section_break_before_paragraph(doc, para_index):
    """
    在指定段落前插入分节符（下一页）
    
    Args:
        doc: Document对象
        para_index: 段落索引
    """
    if para_index <= 0 or para_index >= len(doc.paragraphs):
        return False
    
    # 获取目标段落的前一个段落
    prev_para = doc.paragraphs[para_index - 1]
    
    # 在前一个段落后添加分节符
    pPr = prev_para._element.get_or_add_pPr()
    
    # 创建 sectPr 元素（分节属性）
    sectPr = create_element('w:sectPr')
    
    # 设置分节类型为"下一页"
    sectType = create_element('w:type')
    create_attribute(sectType, 'w:val', 'nextPage')
    sectPr.append(sectType)
    
    # 复制当前节的页面设置
    if doc.sections:
        current_section = doc.sections[0]
        # 页面大小
        pgSz = create_element('w:pgSz')
        create_attribute(pgSz, 'w:w', str(int(current_section.page_width.twips)))
        create_attribute(pgSz, 'w:h', str(int(current_section.page_height.twips)))
        sectPr.append(pgSz)
        
        # 页边距
        pgMar = create_element('w:pgMar')
        create_attribute(pgMar, 'w:top', str(int(current_section.top_margin.twips)))
        create_attribute(pgMar, 'w:right', str(int(current_section.right_margin.twips)))
        create_attribute(pgMar, 'w:bottom', str(int(current_section.bottom_margin.twips)))
        create_attribute(pgMar, 'w:left', str(int(current_section.left_margin.twips)))
        create_attribute(pgMar, 'w:header', '720')
        create_attribute(pgMar, 'w:footer', '720')
        sectPr.append(pgMar)
    
    pPr.append(sectPr)
    return True


def set_page_number_start(section, start_num=1):
    """
    设置节的页码起始值
    
    Args:
        section: Section对象
        start_num: 起始页码
    """
    sectPr = section._sectPr
    
    # 创建或获取 pgNumType 元素
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = create_element('w:pgNumType')
        sectPr.append(pgNumType)
    
    # 设置起始页码
    create_attribute(pgNumType, 'w:start', str(start_num))


def set_footer_content(section, company_name, show_page_number=True):
    """
    设置单个节的页脚内容
    
    Args:
        section: Section对象
        company_name: 公司名称
        show_page_number: 是否显示页码
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # 清空现有内容
    for paragraph in footer.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    if not show_page_number:
        # 不显示页脚内容
        return
    
    # 创建页脚段落
    para = footer.add_paragraph()
    
    # 设置段落制表位
    pPr = para._element.get_or_add_pPr()
    tabs = create_element('w:tabs')
    
    # 居中制表位（页面中央）
    tab_center = create_element('w:tab')
    create_attribute(tab_center, 'w:val', 'center')
    create_attribute(tab_center, 'w:pos', '4680')  # 约8.25cm，A4纸中央
    tabs.append(tab_center)
    
    # 右对齐制表位
    tab_right = create_element('w:tab')
    create_attribute(tab_right, 'w:val', 'right')
    create_attribute(tab_right, 'w:pos', '9360')  # 约16.5cm
    tabs.append(tab_right)
    
    pPr.append(tabs)
    
    # 添加公司名称
    run_company = para.add_run(company_name)
    run_company.font.size = Pt(9)
    run_company.font.name = '宋体'
    run_company._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加制表符
    para.add_run('\t')
    
    # 添加页码域（只显示数字）
    add_page_number(para)


def set_footer(doc, company_name):
    """
    设置页脚：公司名称（左）+ 页码（居中）
    从"前言"开始编号（页码从1开始）
    
    Args:
        doc: Document对象
        company_name: 公司名称
    """
    # 查找"前言"段落
    preface_index = find_preface_paragraph(doc)
    
    if preface_index > 0:
        print(f'ℹ️ 找到"前言"位置: 第 {preface_index + 1} 段')
        
        # 在"前言"前插入分节符
        if insert_section_break_before_paragraph(doc, preface_index):
            print(f'ℹ️ 已在"前言"前插入分节符')
            
            # 重新加载节（因为插入了新节）
            # 第一节：封面/目录等（不显示页脚）
            # 第二节：从前言开始（显示页脚，页码从1开始）
            
            if len(doc.sections) >= 2:
                # 第一节不显示页脚
                set_footer_content(doc.sections[0], company_name, show_page_number=False)
                
                # 从第二节开始显示页脚，页码从1开始
                for i, section in enumerate(doc.sections[1:], 1):
                    set_footer_content(section, company_name, show_page_number=True)
                    if i == 1:
                        set_page_number_start(section, 1)
            else:
                # 如果只有一个节，直接设置
                for section in doc.sections:
                    set_footer_content(section, company_name, show_page_number=True)
        else:
            print(f"⚠️ 无法插入分节符，为所有节设置页脚")
            for section in doc.sections:
                set_footer_content(section, company_name, show_page_number=True)
    else:
        print(f'ℹ️ 未找到"前言"，为所有节设置页脚')
        # 为所有节设置页脚
        for section in doc.sections:
            set_footer_content(section, company_name, show_page_number=True)


def apply_footer(input_file, company_name=DEFAULT_COMPANY_NAME):
    """
    应用页脚设置
    
    Args:
        input_file: 输入文件路径
        company_name: 公司名称
    """
    input_path = Path(input_file)
    
    # 检查文件是否存在
    if not input_path.exists():
        print(f"❌ 错误: 文件不存在: {input_file}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"❌ 错误: 只支持 .docx 文件")
        sys.exit(1)
    
    print(f"🔄 正在处理文件: {input_path.name}")
    
    # 加载文档
    try:
        doc = Document(str(input_path))
    except Exception as e:
        print(f"❌ 错误: 无法打开文件: {e}")
        sys.exit(1)
    
    # 备份原文件
    backup_path = input_path.with_suffix('.docx.backup')
    try:
        import shutil
        shutil.copy2(str(input_path), str(backup_path))
        print(f"ℹ️ 已备份原文件: {backup_path.name}")
    except Exception as e:
        print(f"⚠️ 备份失败: {e}")
    
    # 设置页脚
    print(f"🔄 正在设置页脚...")
    try:
        set_footer(doc, company_name)
        print(f"✅ 页脚设置完成!")
        print(f"   - 公司名称: {company_name}")
        print(f'   - 页码: 从"前言"开始编号（居中）')
    except Exception as e:
        print(f"❌ 页脚设置失败: {e}")
        sys.exit(1)
    
    # 保存文档
    try:
        doc.save(str(input_path))
        print(f"✅ 已保存: {input_path.name}")
        if backup_path.exists():
            print(f"ℹ️ 如需恢复，请使用备份文件: {backup_path.name}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")
        sys.exit(1)


def main():
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext='docx', allow_multiple=False)
    
    if not files:
        print("页脚设置工具")
        print("\n用法: python3 apply_footer.py <input.docx> [公司名称]")
        print("      或在 Finder 中选择 .docx 文件后运行")
        print("\n示例:")
        print("  python3 apply_footer.py document.docx")
        print(f"  python3 apply_footer.py document.docx \"{DEFAULT_COMPANY_NAME}\"")
        print("\n功能:")
        print(f"  - 左侧: 公司名称（默认: {DEFAULT_COMPANY_NAME}）")
        print('  - 居中: 页码（从"前言"开始编号，封面/目录不显示页码）')
        sys.exit(1)
    
    input_file = files[0]
    company_name = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_COMPANY_NAME
    
    apply_footer(input_file, company_name)


if __name__ == "__main__":
    main()

