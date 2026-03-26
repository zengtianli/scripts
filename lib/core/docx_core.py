#!/usr/bin/env python3
"""
Word 文档核心功能模块
提供所有 docx 相关的转换和格式化功能
"""

import contextlib
import re
import shutil
import subprocess
from pathlib import Path

# 尝试导入 python-docx
try:
    from docx import Document
    from docx.shared import Pt  # noqa: F401
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ==================== 转换功能 ====================

def docx_to_md(input_file: Path, output_file: Path | None = None) -> bool:
    """Word 转 Markdown（使用 markitdown）"""
    if output_file is None:
        output_file = input_file.with_suffix('.md')

    if not shutil.which('markitdown'):
        return False

    try:
        result = subprocess.run(
            ['markitdown', str(input_file)],
            capture_output=True, text=True
        )
        if result.returncode == 0:
            output_file.write_text(result.stdout, encoding='utf-8')
            return True
    except Exception:
        pass
    return False


def doc_to_docx(input_file: Path, output_file: Path | None = None) -> bool:
    """旧版 .doc 转 .docx（使用 textutil）"""
    if output_file is None:
        output_file = input_file.with_suffix('.docx')

    try:
        result = subprocess.run([
            'textutil', '-convert', 'docx',
            str(input_file), '-output', str(output_file)
        ], capture_output=True)
        return result.returncode == 0 and output_file.exists()
    except Exception:
        return False


# ==================== 文本格式化 ====================

def format_text(text: str) -> str:
    """
    格式化文本
    - 统一引号为中文引号
    - 英文标点转中文标点
    - 单位转换（平方米→m²）
    """
    if not text:
        return text

    # 引号替换
    text = re.sub(r'"([^"]*)"', r'"\1"', text)
    text = re.sub(r"'([^']*)'", r"'\1'", text)

    # 标点转换
    punct_map = {',': '，', '.': '。', '?': '？', '!': '！', ':': '：', ';': '；'}
    for en, zh in punct_map.items():
        # 只转换不在数字之间的标点
        text = re.sub(rf'(?<!\d){re.escape(en)}(?!\d)', zh, text)

    # 单位转换
    unit_map = {
        '平方米': 'm²', '立方米': 'm³', '平方公里': 'km²',
        'm2': 'm²', 'm3': 'm³', 'km2': 'km²',
    }
    for old, new in unit_map.items():
        text = text.replace(old, new)

    return text


def format_docx_text(input_file: Path, output_file: Path | None = None) -> bool:
    """格式化 Word 文档中的文本"""
    if not HAS_DOCX:
        return False

    if output_file is None:
        output_file = input_file.parent / f"{input_file.stem}_formatted.docx"

    try:
        doc = Document(input_file)

        # 处理段落
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text:
                    run.text = format_text(run.text)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                run.text = format_text(run.text)

        doc.save(output_file)
        return True
    except Exception:
        return False


def format_numbers_font(input_file: Path, output_file: Path | None = None,
                        font_name: str = 'Times New Roman') -> bool:
    """将数字和英文字母设置为指定字体"""
    if not HAS_DOCX:
        return False

    if output_file is None:
        output_file = input_file.parent / f"{input_file.stem}_font.docx"

    try:
        doc = Document(input_file)

        def process_runs(runs):
            for run in runs:
                if run.text and re.search(r'[a-zA-Z0-9]', run.text):
                    run.font.name = font_name

        for para in doc.paragraphs:
            process_runs(para.runs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_runs(para.runs)

        doc.save(output_file)
        return True
    except Exception:
        return False


# ==================== 样式应用 ====================

def apply_footer(input_file: Path, company_name: str = "",
                 output_file: Path | None = None) -> bool:
    """添加页脚（公司名称 + 页码）"""
    if not HAS_DOCX:
        return False

    if output_file is None:
        output_file = input_file.parent / f"{input_file.stem}_footer.docx"

    try:
        doc = Document(input_file)

        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            para.clear()
            if company_name:
                para.add_run(company_name)

        doc.save(output_file)
        return True
    except Exception:
        return False


def apply_header(input_file: Path, header_text: str = "",
                 output_file: Path | None = None) -> bool:
    """添加页眉"""
    if not HAS_DOCX:
        return False

    if output_file is None:
        output_file = input_file.parent / f"{input_file.stem}_header.docx"

    try:
        doc = Document(input_file)

        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False

            if header.paragraphs:
                para = header.paragraphs[0]
            else:
                para = header.add_paragraph()

            para.clear()
            if header_text:
                para.add_run(header_text)

        doc.save(output_file)
        return True
    except Exception:
        return False


def apply_table_style(input_file: Path, style_name: str = "Table Grid",
                      output_file: Path | None = None) -> bool:
    """应用表格样式"""
    if not HAS_DOCX:
        return False

    if output_file is None:
        output_file = input_file.parent / f"{input_file.stem}_table.docx"

    try:
        doc = Document(input_file)

        for table in doc.tables:
            with contextlib.suppress(KeyError):
                table.style = style_name

        doc.save(output_file)
        return True
    except Exception:
        return False


def format_docx_text(input_path: Path) -> bool:  # noqa: F811
    """格式化 docx 文本(引号、标点)"""
    try:
        import re

        from docx import Document
        doc = Document(str(input_path))
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text:
                    t = run.text
                    t = re.sub(r'"([^"]*)"', r'"\1"', t)
                    t = re.sub(r"'([^']*)'", r"'\1'", t)
                    run.text = t
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                t = run.text
                                t = re.sub(r'"([^"]*)"', r'"\1"', t)
                                run.text = t
        doc.save(str(input_path))
        return True
    except Exception as e:
        print(f"错误: {e}")
        return False

def apply_footer(input_path: Path, text: str) -> bool:  # noqa: F811
    """添加页脚"""
    try:
        from docx import Document
        doc = Document(str(input_path))
        for section in doc.sections:
            footer = section.footer
            if not footer.paragraphs:
                footer.paragraphs.append(doc.add_paragraph())
            footer.paragraphs[0].text = text if text else "页脚"
        doc.save(str(input_path))
        return True
    except Exception as e:
        print(f"错误: {e}")
        return False

def apply_header(input_path: Path, text: str) -> bool:  # noqa: F811
    """添加页眉"""
    try:
        from docx import Document
        doc = Document(str(input_path))
        for section in doc.sections:
            header = section.header
            if not header.paragraphs:
                header.paragraphs.append(doc.add_paragraph())
            header.paragraphs[0].text = text if text else "页眉"
        doc.save(str(input_path))
        return True
    except Exception as e:
        print(f"错误: {e}")
        return False

def apply_table_style(input_path: Path, style_name: str = "ZDWP表格内容") -> bool:  # noqa: F811
    """应用表格样式"""
    try:
        from docx import Document
        doc = Document(str(input_path))
        for table in doc.tables:
            with contextlib.suppress(Exception):
                table.style = style_name
        doc.save(str(input_path))
        return True
    except Exception as e:
        print(f"错误: {e}")
        return False
