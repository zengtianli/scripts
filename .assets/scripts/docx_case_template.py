#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-case-template
# @raycast.mode fullOutput
# @raycast.icon 📝
# @raycast.packageName Scripts
# @raycast.description 创建案例编写格式模板
"""
案例编写格式模板生成工具

根据案例编写体例格式要求，生成带预设样式的 docx 模板：
- 主标题：方正小标宋简体 二号
- 副标题：楷体GB2312 三号
- 摘要/关键词/引言标题：黑体 三号
- 摘要/关键词/引言内容：楷体GB2312 三号
- 一级标题：黑体 三号
- 二级标题：楷体GB2312 三号加粗
- 三级标题：仿宋GB2312 三号加粗
- 正文：仿宋GB2312 三号

用法:
    python3 docx_case_template.py                    # 创建新模板
    python3 docx_case_template.py <existing.docx>   # 对已有文件应用样式
    python3 docx_case_template.py -o output.docx    # 指定输出路径

示例:
    python3 docx_case_template.py
    python3 docx_case_template.py 我的案例.docx
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# ==================== 字号定义 ====================
# 中文字号对应磅值
FONT_SIZES = {
    '初号': 42,
    '小初': 36,
    '一号': 26,
    '小一': 24,
    '二号': 22,
    '小二': 18,
    '三号': 16,
    '小三': 15,
    '四号': 14,
    '小四': 12,
    '五号': 10.5,
    '小五': 9,
}

# ==================== 样式配置 ====================
# 格式：(样式名, 字体, 字号pt, 是否加粗, 描述)
STYLE_CONFIG = {
    # 内置样式修改
    'Title': ('方正小标宋简体', 22, False, '主标题'),
    'Subtitle': ('楷体_GB2312', 16, False, '副标题'),
    'Heading 1': ('黑体', 16, False, '一级标题'),
    'Heading 2': ('楷体_GB2312', 16, True, '二级标题'),
    'Heading 3': ('仿宋_GB2312', 16, True, '三级标题'),
    'Normal': ('仿宋_GB2312', 16, False, '正文'),
    
    # 自定义样式
    '摘要标题': ('黑体', 16, False, '摘要/关键词/引言标题'),
    '摘要内容': ('楷体_GB2312', 16, False, '摘要/关键词/引言内容'),
}

# 字体回退（如果首选字体不存在）
FONT_FALLBACK = {
    '方正小标宋简体': ['FZXiaoBiaoSong-B05S', '华文中宋', 'STZhongsong', '宋体'],
    '楷体_GB2312': ['楷体', 'KaiTi', 'KaiTi_GB2312', '华文楷体'],
    '黑体': ['SimHei', '华文黑体', 'STHeiti'],
    '仿宋_GB2312': ['仿宋', 'FangSong', 'FangSong_GB2312', '华文仿宋'],
}


def set_chinese_font(run_or_font, font_name):
    """
    设置中文字体（同时设置西文和中文字体）
    
    Args:
        run_or_font: Run 对象或 Font 对象
        font_name: 字体名称
    """
    # 获取 font 对象
    if hasattr(run_or_font, 'font'):
        font = run_or_font.font
    else:
        font = run_or_font
    
    # 设置西文字体
    font.name = font_name
    
    # 设置中文字体（关键！）
    font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def setup_style(doc, style_name, font_name, font_size_pt, bold=False):
    """
    设置或创建样式
    
    Args:
        doc: Document 对象
        style_name: 样式名称
        font_name: 字体名称
        font_size_pt: 字号（磅值）
        bold: 是否加粗
    
    Returns:
        Style 对象
    """
    styles = doc.styles
    
    # 尝试获取已有样式
    try:
        style = styles[style_name]
    except KeyError:
        # 创建新样式（基于 Normal）
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
    
    # 设置字体
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)
    font.bold = bold
    
    # 清除颜色（设为自动/黑色）
    font.color.rgb = RGBColor(0, 0, 0)
    
    # 清除其他可能的格式
    font.italic = False
    font.underline = False
    font.strike = False
    font.shadow = False
    font.outline = False
    
    # 设置中文字体
    font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    return style


def create_template(output_path=None):
    """
    创建带预设样式的 docx 模板
    
    Args:
        output_path: 输出路径，默认为当前目录下的 "案例模板.docx"
    
    Returns:
        Path: 生成的文件路径
    """
    if output_path is None:
        output_path = Path.cwd() / '案例模板.docx'
    else:
        output_path = Path(output_path)
    
    print("📝 正在创建案例编写格式模板...")
    
    # 创建新文档
    doc = Document()
    
    # 设置所有样式
    print("🔧 配置样式...")
    for style_name, (font_name, font_size, bold, desc) in STYLE_CONFIG.items():
        style = setup_style(doc, style_name, font_name, font_size, bold)
        print(f"   ✓ {style_name}: {font_name} {font_size}pt {'加粗' if bold else ''} ({desc})")
    
    # 添加示例内容（展示各样式效果）
    print("📄 添加示例内容...")
    
    # 主标题
    p = doc.add_paragraph('案例主标题（方正小标宋简体 二号）', style='Title')
    
    # 副标题
    p = doc.add_paragraph('——副标题（楷体GB2312 三号）', style='Subtitle')
    
    # 摘要
    p = doc.add_paragraph('摘要', style='摘要标题')
    p = doc.add_paragraph('这是摘要内容示例，使用楷体GB2312三号字体。要高度概括案例发生的背景情况、面临的问题和矛盾、主要做法和成效、经验启示等，字数在300字以内。', style='摘要内容')
    
    # 关键词
    p = doc.add_paragraph('关键词', style='摘要标题')
    p = doc.add_paragraph('关键词1；关键词2；关键词3', style='摘要内容')
    
    # 引言
    p = doc.add_paragraph('引言', style='摘要标题')
    p = doc.add_paragraph('这是引言内容示例，使用楷体GB2312三号字体。要概括写作意图，说明目的和意义，指出写作范围。字数在300字以内。', style='摘要内容')
    
    # 正文标题层级
    p = doc.add_paragraph('一、一级标题（黑体 三号）', style='Heading 1')
    p = doc.add_paragraph('这是正文内容，使用仿宋GB2312三号字体。', style='Normal')
    
    p = doc.add_paragraph('（一）二级标题（楷体GB2312 三号加粗）', style='Heading 2')
    p = doc.add_paragraph('这是正文内容，使用仿宋GB2312三号字体。要素齐全、资料详实、观点鲜明、逻辑清晰，文字流畅、通俗易懂。', style='Normal')
    
    p = doc.add_paragraph('1. 三级标题（仿宋GB2312 三号加粗）', style='Heading 3')
    p = doc.add_paragraph('这是正文内容，使用仿宋GB2312三号字体。', style='Normal')
    
    # 添加样式说明表格
    doc.add_paragraph()  # 空行
    p = doc.add_paragraph('【样式对照表】', style='Heading 1')
    
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['元素', 'Word样式名', '字体', '字号']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # 数据行
    data = [
        ('主标题', 'Title', '方正小标宋简体', '二号(22pt)'),
        ('副标题', 'Subtitle', '楷体GB2312', '三号(16pt)'),
        ('摘要/关键词/引言标题', '摘要标题', '黑体', '三号(16pt)'),
        ('摘要/关键词/引言内容', '摘要内容', '楷体GB2312', '三号(16pt)'),
        ('一级标题', 'Heading 1', '黑体', '三号(16pt)'),
        ('二级标题', 'Heading 2', '楷体GB2312', '三号(16pt)加粗'),
        ('三级标题', 'Heading 3', '仿宋GB2312', '三号(16pt)加粗'),
        ('正文', 'Normal', '仿宋GB2312', '三号(16pt)'),
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    # 保存文档
    doc.save(str(output_path))
    print(f"✅ 模板创建成功: {output_path}")
    
    return output_path


def apply_styles_to_existing(input_path, output_path=None):
    """
    将样式应用到已有文档（只修改样式定义，不改变内容的样式引用）
    
    Args:
        input_path: 输入文件路径
        output_path: 输出路径，默认覆盖原文件
    
    Returns:
        Path: 输出文件路径
    """
    input_path = Path(input_path)
    
    if not input_path.exists():
        print(f"❌ 文件不存在: {input_path}")
        sys.exit(1)
    
    if output_path is None:
        output_path = input_path
    else:
        output_path = Path(output_path)
    
    print(f"📝 正在为文档应用案例编写格式样式: {input_path.name}")
    
    # 打开文档
    doc = Document(str(input_path))
    
    # 备份（如果是覆盖原文件）
    if output_path == input_path:
        backup_path = input_path.with_suffix('.docx.backup')
        import shutil
        shutil.copy2(str(input_path), str(backup_path))
        print(f"ℹ️ 已备份原文件: {backup_path.name}")
    
    # 设置所有样式
    print("🔧 配置样式...")
    for style_name, (font_name, font_size, bold, desc) in STYLE_CONFIG.items():
        style = setup_style(doc, style_name, font_name, font_size, bold)
        print(f"   ✓ {style_name}: {font_name} {font_size}pt {'加粗' if bold else ''}")
    
    # 保存文档
    doc.save(str(output_path))
    print(f"✅ 样式应用成功: {output_path}")
    
    return output_path


def main():
    args = sys.argv[1:]
    
    output_path = None
    input_path = None
    
    # 解析参数
    i = 0
    while i < len(args):
        if args[i] in ['-o', '--output']:
            if i + 1 < len(args):
                output_path = args[i + 1]
                i += 2
            else:
                print("❌ -o 参数需要指定输出路径")
                sys.exit(1)
        elif args[i] in ['-h', '--help']:
            print(__doc__)
            sys.exit(0)
        elif not args[i].startswith('-'):
            input_path = args[i]
            i += 1
        else:
            print(f"❌ 未知参数: {args[i]}")
            sys.exit(1)
    
    if input_path:
        # 对已有文件应用样式
        apply_styles_to_existing(input_path, output_path)
    else:
        # 创建新模板
        create_template(output_path)
    
    print()
    print("📋 使用方法:")
    print("   1. 用 Word 打开生成的文档")
    print("   2. 选中要格式化的内容")
    print("   3. 在「样式」面板中选择对应样式")
    print()
    print("💡 样式快捷键提示:")
    print("   Ctrl+Alt+1 = Heading 1 (一级标题)")
    print("   Ctrl+Alt+2 = Heading 2 (二级标题)")
    print("   Ctrl+Alt+3 = Heading 3 (三级标题)")


if __name__ == "__main__":
    main()

