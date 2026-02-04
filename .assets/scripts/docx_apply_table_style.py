#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-table
# @raycast.mode fullOutput
# @raycast.icon 📄
# @raycast.packageName Scripts
# @raycast.description Apply table style
"""
应用表格样式工具
将文档中所有表格应用指定样式，并设置实线边框

功能：
1. 应用表格内容样式（支持表格样式或段落样式）
2. 设置表格实线边框
3. 表头（第一行）加粗
4. 应用表名样式（"ZDWP表名"）到表格前面的段落
5. 表格后面添加空行

支持两种样式模式：
- 表格样式：应用到表格本身
- 段落样式：应用到表格中所有单元格的段落（如"ZDWP表格内容"）

用法:
    python3 apply_table_style.py <input.docx> [样式名称]
    
示例:
    python3 apply_table_style.py document.docx "ZDWP表格内容"
    python3 apply_table_style.py document.docx  # 默认使用"ZDWP表格内容"
"""

import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

from common_utils import get_input_files

from common_utils import get_input_files

def find_style_fuzzy(doc, style_name):
    """
    模糊查找文档中的样式（表格样式或段落样式）
    支持：
    1. 精确匹配
    2. 忽略空格差异（"ZDWP表格内容" 能匹配 "ZDWP 表格内容"）
    3. 包含关键词（如果精确+空格都不匹配，则查找包含关键词的）
    
    Args:
        doc: Document对象
        style_name: 样式名称
        
    Returns:
        tuple: (实际样式名称或None, 样式类型) - 类型为 'table' 或 'paragraph'
    """
    try:
        styles = doc.styles
        
        # 准备搜索用的名称（去除所有空格）
        search_name_normalized = style_name.replace(' ', '').replace('\u3000', '')
        
        exact_match = None
        space_match = None
        partial_matches = []
        
        for style in styles:
            if not style.name:
                continue
                
            style_type = None
            if style.type == 3:  # TABLE
                style_type = 'table'
            elif style.type == 1:  # PARAGRAPH
                style_type = 'paragraph'
            else:
                continue
            
            # 1. 精确匹配
            if style.name == style_name:
                exact_match = (style.name, style_type)
                break
            
            # 2. 忽略空格的匹配
            style_normalized = style.name.replace(' ', '').replace('\u3000', '')
            if style_normalized == search_name_normalized:
                space_match = (style.name, style_type)
            
            # 3. 包含关键词的匹配（两个方向都检查）
            if search_name_normalized in style_normalized or style_normalized in search_name_normalized:
                # 只记录相关性高的（长度差距不要太大）
                if abs(len(style_normalized) - len(search_name_normalized)) <= 5:
                    partial_matches.append((style.name, style_type))
        
        # 返回优先级：精确匹配 > 空格匹配 > 部分匹配
        if exact_match:
            return exact_match
        if space_match:
            return space_match
        if partial_matches:
            return partial_matches[0]  # 返回第一个部分匹配
        
        return None, None
        
    except Exception as e:
        print(f"⚠️ 样式查找出错: {e}")
        return None, None

def check_style_exists(doc, style_name):
    """
    检查文档中是否存在指定的样式（表格样式或段落样式）
    使用模糊匹配
    
    Args:
        doc: Document对象
        style_name: 样式名称
        
    Returns:
        tuple: (是否存在, 样式类型) - 类型为 'table' 或 'paragraph'
    """
    found_name, style_type = find_style_fuzzy(doc, style_name)
    return (found_name is not None), style_type

def set_table_border_solid(table):
    """
    设置表格所有边框为实线
    
    Args:
        table: Table对象
    """
    # 获取表格属性
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # 获取或创建表格边框
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    
    # 定义实线边框 (single = 实线)
    border_attrs = {
        qn('w:val'): 'single',      # 实线
        qn('w:sz'): '4',             # 边框宽度 (1/8 point)
        qn('w:space'): '0',          # 边框间距
        qn('w:color'): '000000'      # 黑色
    }
    
    # 设置所有边框：上、下、左、右、水平、垂直
    border_names = ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']
    
    for border_name in border_names:
        border_element = tblBorders.find(qn(f'w:{border_name}'))
        if border_element is None:
            border_element = OxmlElement(f'w:{border_name}')
            tblBorders.append(border_element)
        
        # 设置边框属性
        for attr, value in border_attrs.items():
            border_element.set(attr, value)

def set_table_header_bold(table):
    """
    设置表格第一行（表头）为粗体
    
    Args:
        table: Table对象
    """
    if len(table.rows) == 0:
        return
    
    # 获取第一行
    header_row = table.rows[0]
    
    # 遍历第一行的所有单元格
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

def get_table_caption_paragraph(doc, table):
    """
    获取表格前面的段落（表名）
    
    Args:
        doc: Document对象
        table: Table对象
        
    Returns:
        Paragraph对象或None
    """
    # 获取文档的所有元素
    body_elements = doc.element.body
    
    # 找到当前表格的索引
    table_element = table._element
    
    # 遍历body的所有子元素
    prev_paragraph = None
    for element in body_elements:
        if element == table_element:
            # 找到表格了，返回前一个段落
            return prev_paragraph
        
        # 检查是否是段落
        if element.tag == qn('w:p'):
            # 创建Paragraph对象
            from docx.text.paragraph import Paragraph
            prev_paragraph = Paragraph(element, doc)
    
    return None

def add_blank_line_after_table(doc, table):
    """
    在表格后面添加一个空行
    
    Args:
        doc: Document对象
        table: Table对象
    """
    # 获取文档的body元素
    body = doc.element.body
    
    # 找到表格元素的位置
    table_element = table._element
    table_index = list(body).index(table_element)
    
    # 检查表格后面是否已经有空段落
    if table_index + 1 < len(body):
        next_element = body[table_index + 1]
        # 如果后面是段落且为空，不再添加
        if next_element.tag == qn('w:p'):
            from docx.text.paragraph import Paragraph
            next_para = Paragraph(next_element, doc)
            if not next_para.text.strip():
                return  # 已经有空行了
    
    # 创建新的空段落元素
    new_para = OxmlElement('w:p')
    
    # 在表格后面插入空段落
    body.insert(table_index + 1, new_para)

def apply_table_style(input_file, style_name="ZDWP表格内容"):
    """
    应用样式到文档中的所有表格
    
    功能：
    1. 应用表格内容样式（表格样式或段落样式）
    2. 设置实线边框
    3. 表头（第一行）加粗
    4. 应用"ZDWP表名"样式到表格前面的段落
    5. 表格后面添加空行
    
    Args:
        input_file: 输入文件路径
        style_name: 表格内容样式名称（默认"ZDWP表格内容"）
    """
    input_path = Path(input_file)
    
    # 检查文件是否存在
    if not input_path.exists():
        print(f"❌ 错误: 文件不存在: {input_file}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"❌ 错误: 只支持 .docx 文件")
        sys.exit(1)
    
    print(f"🔄 正在处理文件: {input_path.name}")
    
    # 加载文档
    try:
        doc = Document(str(input_path))
    except Exception as e:
        print(f"❌ 错误: 无法打开文件: {e}")
        sys.exit(1)
    
    # 检查表格内容样式是否存在（使用模糊匹配）
    actual_style_name, style_type = find_style_fuzzy(doc, style_name)
    
    if not actual_style_name:
        print(f"❌ 错误: 文档中不存在样式 '{style_name}' 或类似样式")
        print(f"💡 请在Word中先添加该样式，或检查样式名称是否正确")
        sys.exit(1)
    
    if actual_style_name != style_name:
        print(f"ℹ️ 使用样式: {actual_style_name} (匹配: {style_name})")
    
    if style_type == 'table':
        print(f"✅ 找到表格样式: {actual_style_name}")
    elif style_type == 'paragraph':
        print(f"✅ 找到段落样式: {actual_style_name}")
        print(f"ℹ️ 将应用到表格中的所有单元格段落")
    
    # 使用实际找到的样式名称
    style_name = actual_style_name
    
    # 检查表名样式是否存在（使用模糊匹配）
    caption_style_search = "ZDWP表名"
    actual_caption_style, caption_style_type = find_style_fuzzy(doc, caption_style_search)
    
    if actual_caption_style:
        caption_style_name = actual_caption_style
        if actual_caption_style != caption_style_search:
            print(f"ℹ️ 使用表名样式: {actual_caption_style} (匹配: {caption_style_search})")
        else:
            print(f"✅ 找到表名样式: {actual_caption_style}")
    else:
        print(f"⚠️ 未找到表名样式 '{caption_style_search}' 或类似样式，将跳过表名设置")
        caption_style_name = None
    
    # 获取表格数量
    table_count = len(doc.tables)
    if table_count == 0:
        print("⚠️ 文档中没有表格")
        return
    
    print(f"ℹ️ 文档包含 {table_count} 个表格")
    
    # 备份原文件
    backup_path = input_path.with_suffix('.docx.backup')
    try:
        import shutil
        shutil.copy2(str(input_path), str(backup_path))
        print(f"ℹ️ 已备份原文件: {backup_path.name}")
    except Exception as e:
        print(f"⚠️ 备份失败: {e}")
    
    # 应用样式到所有表格
    success_count = 0
    error_count = 0
    caption_count = 0
    
    print(f"🔄 正在应用样式...")
    
    for i, table in enumerate(doc.tables, 1):
        try:
            # 1. 应用表格内容样式
            if style_type == 'table':
                # 表格样式：应用到表格本身
                table.style = style_name
            elif style_type == 'paragraph':
                # 段落样式：应用到表格中的所有单元格段落
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style = style_name
            
            # 2. 设置实线边框
            set_table_border_solid(table)
            
            # 3. 设置表头（第一行）加粗
            set_table_header_bold(table)
            
            # 4. 应用表名样式（如果存在）
            if caption_style_name:
                caption_para = get_table_caption_paragraph(doc, table)
                if caption_para and caption_para.text.strip():
                    # 只有段落有内容才应用样式
                    caption_para.style = caption_style_name
                    caption_count += 1
            
            # 5. 表格后面添加空行
            add_blank_line_after_table(doc, table)
            
            success_count += 1
            
        except Exception as e:
            print(f"⚠️ 表格 {i} 处理失败: {e}")
            error_count += 1
    
    # 保存文档
    try:
        doc.save(str(input_path))
        print(f"✅ 样式应用完成!")
        print(f"   - 成功处理: {success_count} 个表格")
        print(f"   - 表头加粗: {success_count} 个表格")
        if caption_style_name:
            print(f"   - 表名样式: {caption_count} 个段落")
        if error_count > 0:
            print(f"   - 失败: {error_count} 个表格")
        print(f"   - 已保存: {input_path.name}")
        if backup_path.exists():
            print(f"ℹ️ 如需恢复，请使用备份文件: {backup_path.name}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")
        sys.exit(1)

def main():
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext='docx', allow_multiple=False)
    
    if not files:
        print("用法: python3 apply_table_style.py <input.docx> [样式名称]")
        print("      或在 Finder 中选择 .docx 文件后运行")
        print("示例: python3 apply_table_style.py document.docx")
        print("      python3 apply_table_style.py document.docx \"ZDWP表格内容\"")
        sys.exit(1)
    
    input_file = files[0]
    style_name = sys.argv[2] if len(sys.argv) > 2 else "ZDWP表格内容"
    
    apply_table_style(input_file, style_name)

if __name__ == "__main__":
    main()

