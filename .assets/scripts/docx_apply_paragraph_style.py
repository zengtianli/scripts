#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-paragraph
# @raycast.mode fullOutput
# @raycast.icon 📄
# @raycast.packageName Scripts
# @raycast.description Apply paragraph style
"""
应用段落样式工具
将文档中的正文段落统一应用指定样式

功能：
1. 应用"ZDWP正文"样式到所有正文段落
2. 智能识别：跳过标题、表名、表格内段落等
3. 保留原有格式和内容

用法:
    python3 apply_paragraph_style.py <input.docx> [样式名称]
    
示例:
    python3 apply_paragraph_style.py document.docx "ZDWP正文"
    python3 apply_paragraph_style.py document.docx  # 默认使用"ZDWP正文"
"""

import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from common_utils import get_input_files

from common_utils import get_input_files

def find_style_fuzzy(doc, style_name):
    """
    模糊查找文档中的段落样式
    支持：
    1. 精确匹配
    2. 忽略空格差异（"ZDWP正文" 能匹配 "ZDWP 正文"）
    3. 包含关键词（如果精确+空格都不匹配，则查找包含关键词的）
    
    Args:
        doc: Document对象
        style_name: 样式名称
        
    Returns:
        str或None: 找到的实际样式名称，未找到返回None
    """
    try:
        styles = doc.styles
        
        # 准备搜索用的名称（去除所有空格）
        search_name_normalized = style_name.replace(' ', '').replace('\u3000', '')
        
        exact_match = None
        space_match = None
        partial_matches = []
        
        for style in styles:
            if not style.name or style.type != 1:  # 只查找段落样式
                continue
            
            # 1. 精确匹配
            if style.name == style_name:
                exact_match = style.name
                break
            
            # 2. 忽略空格的匹配
            style_normalized = style.name.replace(' ', '').replace('\u3000', '')
            if style_normalized == search_name_normalized:
                space_match = style.name
            
            # 3. 包含关键词的匹配（两个方向都检查）
            if search_name_normalized in style_normalized or style_normalized in search_name_normalized:
                # 只记录相关性高的（长度差距不要太大）
                if abs(len(style_normalized) - len(search_name_normalized)) <= 5:
                    partial_matches.append(style.name)
        
        # 返回优先级：精确匹配 > 空格匹配 > 部分匹配
        if exact_match:
            return exact_match
        if space_match:
            return space_match
        if partial_matches:
            return partial_matches[0]  # 返回第一个部分匹配
        
        return None
        
    except Exception as e:
        print(f"⚠️ 样式查找出错: {e}")
        return None

def check_style_exists(doc, style_name):
    """
    检查文档中是否存在指定的段落样式
    使用模糊匹配
    
    Args:
        doc: Document对象
        style_name: 样式名称
        
    Returns:
        bool: 样式是否存在
    """
    found_name = find_style_fuzzy(doc, style_name)
    return found_name is not None

def is_heading_paragraph(paragraph):
    """
    判断段落是否是标题
    
    Args:
        paragraph: Paragraph对象
        
    Returns:
        bool: 是否是标题
    """
    if not paragraph.style:
        return False
    
    style_name = paragraph.style.name
    
    # 判断是否是标题样式
    if 'Heading' in style_name or '标题' in style_name or 'Title' in style_name:
        return True
    
    return False

def is_special_paragraph(paragraph):
    """
    判断段落是否是特殊段落（表名、图片说明等）
    
    Args:
        paragraph: Paragraph对象
        
    Returns:
        bool: 是否是特殊段落
    """
    if not paragraph.style:
        return False
    
    style_name = paragraph.style.name
    
    # 判断是否是特殊样式
    special_keywords = ['表名', 'Caption', '图', '表', 'TOC', '目录', 'Header', 'Footer', '页眉', '页脚']
    
    for keyword in special_keywords:
        if keyword in style_name:
            return True
    
    return False

def is_in_table(paragraph):
    """
    判断段落是否在表格内
    
    Args:
        paragraph: Paragraph对象
        
    Returns:
        bool: 是否在表格内
    """
    # 获取段落的父元素
    parent = paragraph._element.getparent()
    
    # 向上遍历，查找是否有表格单元格(tc)
    while parent is not None:
        if parent.tag == qn('w:tc'):  # w:tc = table cell
            return True
        parent = parent.getparent()
    
    return False

def should_apply_style(paragraph):
    """
    判断段落是否应该应用正文样式
    
    Args:
        paragraph: Paragraph对象
        
    Returns:
        bool: 是否应该应用样式
    """
    # 空段落不处理
    if not paragraph.text.strip():
        return False
    
    # 标题不处理
    if is_heading_paragraph(paragraph):
        return False
    
    # 特殊段落（表名等）不处理
    if is_special_paragraph(paragraph):
        return False
    
    # 表格内段落不处理
    if is_in_table(paragraph):
        return False
    
    return True

def apply_paragraph_style(input_file, style_name="ZDWP正文"):
    """
    应用正文样式到文档中的段落
    
    Args:
        input_file: 输入文件路径
        style_name: 段落样式名称（默认"ZDWP正文"）
    """
    input_path = Path(input_file)
    
    # 检查文件是否存在
    if not input_path.exists():
        print(f"❌ 错误: 文件不存在: {input_file}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"❌ 错误: 只支持 .docx 文件")
        sys.exit(1)
    
    print(f"🔄 正在处理文件: {input_path.name}")
    
    # 加载文档
    try:
        doc = Document(str(input_path))
    except Exception as e:
        print(f"❌ 错误: 无法打开文件: {e}")
        sys.exit(1)
    
    # 检查样式是否存在（使用模糊匹配）
    actual_style_name = find_style_fuzzy(doc, style_name)
    
    if not actual_style_name:
        print(f"❌ 错误: 文档中不存在段落样式 '{style_name}' 或类似样式")
        print(f"💡 请在Word中先添加该样式，或检查样式名称是否正确")
        sys.exit(1)
    
    if actual_style_name != style_name:
        print(f"ℹ️ 使用样式: {actual_style_name} (匹配: {style_name})")
    else:
        print(f"✅ 找到段落样式: {actual_style_name}")
    
    # 使用实际找到的样式名称
    style_name = actual_style_name
    
    # 获取段落数量
    paragraph_count = len(doc.paragraphs)
    print(f"ℹ️ 文档包含 {paragraph_count} 个段落")
    
    # 备份原文件
    backup_path = input_path.with_suffix('.docx.backup')
    try:
        import shutil
        shutil.copy2(str(input_path), str(backup_path))
        print(f"ℹ️ 已备份原文件: {backup_path.name}")
    except Exception as e:
        print(f"⚠️ 备份失败: {e}")
    
    # 应用样式到正文段落
    success_count = 0
    skipped_count = 0
    error_count = 0
    
    print(f"🔄 正在应用正文样式...")
    
    for i, paragraph in enumerate(doc.paragraphs, 1):
        try:
            if should_apply_style(paragraph):
                # 应用样式
                paragraph.style = style_name
                success_count += 1
            else:
                skipped_count += 1
                
        except Exception as e:
            print(f"⚠️ 段落 {i} 处理失败: {e}")
            error_count += 1
    
    # 保存文档
    try:
        doc.save(str(input_path))
        print(f"✅ 样式应用完成!")
        print(f"   - 应用样式: {success_count} 个段落")
        print(f"   - 跳过: {skipped_count} 个段落（标题、表名、表格内等）")
        if error_count > 0:
            print(f"   - 失败: {error_count} 个段落")
        print(f"   - 已保存: {input_path.name}")
        if backup_path.exists():
            print(f"ℹ️ 如需恢复，请使用备份文件: {backup_path.name}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")
        sys.exit(1)

def main():
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext='docx', allow_multiple=False)
    
    if not files:
        print("用法: python3 apply_paragraph_style.py <input.docx> [样式名称]")
        print("      或在 Finder 中选择 .docx 文件后运行")
        print("示例: python3 apply_paragraph_style.py document.docx")
        print("      python3 apply_paragraph_style.py document.docx \"ZDWP正文\"")
        sys.exit(1)
    
    input_file = files[0]
    style_name = sys.argv[2] if len(sys.argv) > 2 else "ZDWP正文"
    
    apply_paragraph_style(input_file, style_name)

if __name__ == "__main__":
    main()

