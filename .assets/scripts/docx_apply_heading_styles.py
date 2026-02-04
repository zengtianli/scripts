#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-heading
# @raycast.mode fullOutput
# @raycast.icon 📄
# @raycast.packageName Scripts
# @raycast.description Apply heading styles
"""
自动识别并应用标题样式工具
根据段落开头的编号自动识别标题层级并应用对应的标题样式

识别规则：
- 1, 2, 3... → 标题 1
- 前  言, 附  录（带空格）→ 标题 1
- 1.1, 1.1编制目的... → 标题 2（空格可选）
- 1.1.1, 1.1.1主要目的... → 标题 3（空格可选）
- 1.1.1.1, 1.1.1.1技术规范... → 标题 4（空格可选）
- （1）、（2）、（3）... → 标题 4

用法:
    python3 apply_heading_styles.py <input.docx>
    
示例:
    python3 apply_heading_styles.py document.docx
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from common_utils import get_input_files

# 标题识别的正则表达式（按优先级排序：从特殊到一般）
HEADING_PATTERNS = [
    # 特殊格式：带空格的标题（如"前  言"）
    (r'^[\u4e00-\u9fa5]\s+[\u4e00-\u9fa5]$', 'Heading 1', 1),  # 前  言（单字+空格+单字）
    
    # 带括号的编号
    (r'^[（(]\d+[)）]\s*', 'Heading 4', 4),          # （1）或(1)
    
    # 数字编号（空格可选，从最长到最短）
    (r'^\d+\.\d+\.\d+\.\d+\s*', 'Heading 4', 4),  # 1.1.1.1 或 1.1.1.1编制目的
    (r'^\d+\.\d+\.\d+\s*', 'Heading 3', 3),        # 1.1.1 或 1.1.1主要目的
    (r'^\d+\.\d+\s*', 'Heading 2', 2),             # 1.1 或 1.1编制目的
    (r'^\d+\s+', 'Heading 1', 1),                  # 1 总论（保留空格要求，避免误匹配）
]

def is_in_table(paragraph):
    """
    判断段落是否在表格内
    
    Args:
        paragraph: Paragraph对象
        
    Returns:
        bool: 是否在表格内
    """
    parent = paragraph._element.getparent()
    
    while parent is not None:
        if parent.tag == qn('w:tc'):  # w:tc = table cell
            return True
        parent = parent.getparent()
    
    return False

def detect_heading_level(text):
    """
    检测段落文本的标题层级
    
    Args:
        text: 段落文本
        
    Returns:
        tuple: (标题样式名称, 层级) 或 (None, 0)
    """
    text = text.strip()
    
    # 按顺序匹配（从最长的开始，避免误匹配）
    for pattern, style_name, level in HEADING_PATTERNS:
        if re.match(pattern, text):
            return style_name, level
    
    return None, 0

def is_already_heading(paragraph):
    """
    判断段落是否已经是标题样式
    
    Args:
        paragraph: Paragraph对象
        
    Returns:
        bool: 是否已经是标题
    """
    if not paragraph.style:
        return False
    
    style_name = paragraph.style.name
    
    # 检查是否已经是标题样式
    if 'Heading' in style_name or '标题' in style_name:
        return True
    
    return False

def apply_heading_styles(input_file):
    """
    自动识别并应用标题样式
    
    Args:
        input_file: 输入文件路径
    """
    input_path = Path(input_file)
    
    # 检查文件是否存在
    if not input_path.exists():
        print(f"❌ 错误: 文件不存在: {input_file}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"❌ 错误: 只支持 .docx 文件")
        sys.exit(1)
    
    print(f"🔄 正在处理文件: {input_path.name}")
    
    # 加载文档
    try:
        doc = Document(str(input_path))
    except Exception as e:
        print(f"❌ 错误: 无法打开文件: {e}")
        sys.exit(1)
    
    # 获取段落数量
    paragraph_count = len(doc.paragraphs)
    print(f"ℹ️ 文档包含 {paragraph_count} 个段落")
    
    # 备份原文件
    backup_path = input_path.with_suffix('.docx.backup')
    try:
        import shutil
        shutil.copy2(str(input_path), str(backup_path))
        print(f"ℹ️ 已备份原文件: {backup_path.name}")
    except Exception as e:
        print(f"⚠️ 备份失败: {e}")
    
    # 统计信息
    heading_counts = {1: 0, 2: 0, 3: 0, 4: 0}
    skipped_count = 0
    error_count = 0
    
    print(f"🔄 正在识别并应用标题样式...")
    
    for i, paragraph in enumerate(doc.paragraphs, 1):
        try:
            # 空段落跳过
            if not paragraph.text.strip():
                skipped_count += 1
                continue
            
            # 表格内段落跳过
            if is_in_table(paragraph):
                skipped_count += 1
                continue
            
            # 已经是标题样式的跳过
            if is_already_heading(paragraph):
                skipped_count += 1
                continue
            
            # 检测标题层级
            style_name, level = detect_heading_level(paragraph.text)
            
            if style_name:
                # 应用标题样式
                paragraph.style = style_name
                heading_counts[level] += 1
            else:
                skipped_count += 1
                
        except Exception as e:
            print(f"⚠️ 段落 {i} 处理失败: {e}")
            error_count += 1
    
    # 保存文档
    try:
        doc.save(str(input_path))
        print(f"✅ 标题样式应用完成!")
        print(f"   - 标题 1: {heading_counts[1]} 个")
        print(f"   - 标题 2: {heading_counts[2]} 个")
        print(f"   - 标题 3: {heading_counts[3]} 个")
        print(f"   - 标题 4: {heading_counts[4]} 个")
        print(f"   - 跳过: {skipped_count} 个段落")
        if error_count > 0:
            print(f"   - 失败: {error_count} 个段落")
        print(f"   - 已保存: {input_path.name}")
        if backup_path.exists():
            print(f"ℹ️ 如需恢复，请使用备份文件: {backup_path.name}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")
        sys.exit(1)

def main():
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext='docx', allow_multiple=False)
    
    if not files:
        print("用法: python3 apply_heading_styles.py <input.docx>")
        print("      或在 Finder 中选择 .docx 文件后运行")
        print("示例: python3 apply_heading_styles.py document.docx")
        sys.exit(1)
    
    input_file = files[0]
    
    apply_heading_styles(input_file)

if __name__ == "__main__":
    main()

