#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从现有 DOCX 创建空白模板（保留样式）

使用 python-docx 确保模板结构完整
"""

import sys
import os
from docx import Document
from docx.shared import Pt, Cm


def create_template_from_docx(source_docx, output_path):
    """从现有 docx 创建空白模板"""
    from docx.oxml.ns import qn
    from lxml import etree
    
    # 加载源文档
    doc = Document(source_docx)
    
    # 找到并保存最后一个 sectPr
    body = doc.element.body
    sect_pr = None
    for element in body[:]:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'sectPr':
            sect_pr = element
    
    # 清空所有内容
    for element in body[:]:
        body.remove(element)
    
    # 添加一个空段落
    doc.add_paragraph()
    
    # 恢复 sectPr（必须在最后）
    if sect_pr is not None:
        body.append(sect_pr)
    else:
        # 如果没有，创建默认的
        from docx.oxml import OxmlElement
        sect_pr = OxmlElement('w:sectPr')
        pgSz = OxmlElement('w:pgSz')
        pgSz.set(qn('w:w'), '11906')
        pgSz.set(qn('w:h'), '16838')
        sect_pr.append(pgSz)
        pgMar = OxmlElement('w:pgMar')
        pgMar.set(qn('w:top'), '1440')
        pgMar.set(qn('w:right'), '1800')
        pgMar.set(qn('w:bottom'), '1440')
        pgMar.set(qn('w:left'), '1800')
        sect_pr.append(pgMar)
        body.append(sect_pr)
    
    # 保存
    doc.save(output_path)
    print(f"✅ 模板已创建: {output_path}")
    print(f"   源文件: {source_docx}")
    print(f"   样式数量: {len(doc.styles)}")


def create_zdwp_template_manual(output_path):
    """手动创建带有 ZDWP 样式的模板"""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    
    doc = Document()
    
    # 创建 ZDWP正文 样式
    try:
        zdwp_style = doc.styles.add_style('ZDWP正文', WD_STYLE_TYPE.PARAGRAPH)
        zdwp_style.base_style = doc.styles['Normal']
        
        # 段落格式
        pf = zdwp_style.paragraph_format
        pf.first_line_indent = Cm(0.97)
        pf.line_spacing = 1.5
        pf.alignment = 3  # JUSTIFY
        
        # 字体格式
        font = zdwp_style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
    except Exception as e:
        print(f"创建样式时出错: {e}")
    
    # 添加空段落
    doc.add_paragraph()
    
    doc.save(output_path)
    print(f"✅ 模板已创建: {output_path}")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python create_zdwp_template.py <source.docx> [output.docx]")
        print("      python create_zdwp_template.py --manual [output.docx]")
        sys.exit(1)
    
    if sys.argv[1] == '--manual':
        output = sys.argv[2] if len(sys.argv) > 2 else 'zdwp_template.docx'
        create_zdwp_template_manual(output)
    else:
        source = sys.argv[1]
        output = sys.argv[2] if len(sys.argv) > 2 else 'zdwp_template.docx'
        create_template_from_docx(source, output)

