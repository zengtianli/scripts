#!/usr/bin/env python3
"""
md_docx_heading_template.py - Markdown 转 Docx（样式复刻版）

功能：
  1. 从模板 docx 提取 heading 1-4 + 正文样式
  2. 把 Markdown 转成带样式的 docx

用法：
  # 提取样式
  python md_docx_heading_template.py extract 模板.docx

  # 转换（用提取的样式）
  python md_docx_heading_template.py convert input.md -o output.docx

  # 一步完成
  python md_docx_heading_template.py input.md -t 模板.docx -o output.docx
"""

import argparse
import json
import os
import re
import sys
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent.parent / "lib"))

try:
    from lxml import etree
except ImportError:
    print("❌ 需要安装 lxml: pip install lxml")
    sys.exit(1)

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import nsmap, qn
    from docx.shared import Cm, Pt, Twips
except ImportError:
    print("❌ 需要安装 python-docx: pip install python-docx")
    sys.exit(1)


# Word XML 命名空间
from docx_xml import NSMAP

# 需要提取的样式 ID
TARGET_STYLES = {
    "a": "Normal",  # 基础样式
    "10": "Heading 1",  # 标题1
    "2": "Heading 2",  # 标题2
    "3": "Heading 3",  # 标题3
    "4": "Heading 4",  # 标题4
    "ZDWP": "ZDWP正文",  # 水利正文
    "ZDWP1": "ZDWP表名",  # 表格标题
    "ZDWP3": "ZDWP表格内容",  # 表格单元格文字
    "ZDWP4": "ZDWP图名",  # 图片标题
}

# 样式依赖关系
STYLE_DEPS = {
    "10": ["2"],  # Heading 1 依赖 Heading 2
    "2": ["a"],  # Heading 2 依赖 Normal
    "3": ["4"],  # Heading 3 依赖 Heading 4
    "4": ["a"],  # Heading 4 依赖 Normal
    "ZDWP": ["a"],  # ZDWP正文 依赖 Normal
    "ZDWP1": ["a"],  # ZDWP表名 依赖 Normal
    "ZDWP3": ["a"],  # ZDWP表格内容 依赖 Normal
    "ZDWP4": ["a"],  # ZDWP图名 依赖 Normal
}

# 默认配置文件路径
DEFAULT_STYLES_DIR = os.path.dirname(os.path.abspath(__file__))


def extract_styles_xml(docx_path, output_dir=None):
    """从 docx 提取样式 XML"""

    if output_dir is None:
        output_dir = os.path.dirname(docx_path) or "."

    print(f"📄 提取样式: {docx_path}")

    # 解压 docx
    with zipfile.ZipFile(docx_path, "r") as zf:
        styles_xml = zf.read("word/styles.xml")

    # 解析 XML
    root = etree.fromstring(styles_xml)

    # 收集需要的样式
    collected_styles = {}
    found_body_style = None

    for style in root.findall(".//w:style", NSMAP):
        style_id = style.get(f"{{{NSMAP['w']}}}styleId")
        style_type = style.get(f"{{{NSMAP['w']}}}type")

        # 只要段落样式
        if style_type != "paragraph":
            continue

        # 获取样式名
        name_elem = style.find("w:name", NSMAP)
        style_name = name_elem.get(f"{{{NSMAP['w']}}}val") if name_elem is not None else ""

        # 检查是否是目标样式
        if style_id in TARGET_STYLES:
            collected_styles[style_id] = {"element": deepcopy(style), "name": style_name, "id": style_id}
            print(f"  ✓ 找到: {style_id} ({style_name})")

            # 记录正文样式
            if style_id == "ZDWP":
                found_body_style = "ZDWP"

        # 如果没有 ZDWP，用 Normal 作为正文
        if style_id == "a" and found_body_style is None:
            found_body_style = "a"

    # 检查必要样式
    required = ["a", "10", "2", "3", "4"]
    missing = [s for s in required if s not in collected_styles]
    if missing:
        print(f"⚠️  缺少样式: {missing}")

    # 提取 docDefaults（默认字体设置）
    doc_defaults = root.find(".//w:docDefaults", NSMAP)

    # 生成精简的 styles.xml
    new_root = etree.Element(f"{{{NSMAP['w']}}}styles", nsmap=NSMAP)

    # 添加 docDefaults
    if doc_defaults is not None:
        new_root.append(deepcopy(doc_defaults))

    # 按依赖顺序添加样式
    added = set()

    def add_style(style_id):
        if style_id in added or style_id not in collected_styles:
            return
        # 先添加依赖
        for dep in STYLE_DEPS.get(style_id, []):
            add_style(dep)
        new_root.append(collected_styles[style_id]["element"])
        added.add(style_id)

    for style_id in collected_styles:
        add_style(style_id)

    # 保存 XML
    styles_xml_path = os.path.join(output_dir, "heading_styles.xml")
    tree = etree.ElementTree(new_root)
    tree.write(styles_xml_path, encoding="UTF-8", xml_declaration=True, pretty_print=True)
    print(f"💾 样式 XML: {styles_xml_path}")

    # 生成人类可读的说明
    info_path = os.path.join(output_dir, "styles_info.txt")
    with open(info_path, "w", encoding="utf-8") as f:
        f.write("=" * 60 + "\n")
        f.write(f"样式提取自: {os.path.basename(docx_path)}\n")
        f.write("=" * 60 + "\n\n")

        f.write("提取的样式:\n")
        for style_id, info in collected_styles.items():
            f.write(f"  - {info['id']}: {info['name']}\n")

        f.write(f"\n正文样式: {found_body_style}\n")

        f.write("\nMarkdown 映射:\n")
        f.write("  # 标题    → Heading 1\n")
        f.write("  ## 标题   → Heading 2\n")
        f.write("  ### 标题  → Heading 3\n")
        f.write("  #### 标题 → Heading 4\n")
        f.write(f"  普通段落  → {collected_styles.get(found_body_style, {}).get('name', 'Normal')}\n")
        f.write(f"  表x ...   → {collected_styles.get('ZDWP1', {}).get('name', 'Normal')}\n")
        f.write(f"  表格内容  → {collected_styles.get('ZDWP3', {}).get('name', 'Normal')}\n")
        f.write(f"  图x ...   → {collected_styles.get('ZDWP4', {}).get('name', 'Normal')}\n")

    print(f"📝 样式说明: {info_path}")

    # 保存配置
    config = {
        "source": os.path.basename(docx_path),
        "body_style": found_body_style,
        "body_style_name": collected_styles.get(found_body_style, {}).get("name", "Normal"),
        "styles": {k: v["name"] for k, v in collected_styles.items()},
    }
    config_path = os.path.join(output_dir, "styles_config.json")
    with open(config_path, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

    return styles_xml_path, config


def create_docx_with_styles(styles_xml_path, output_path):
    """创建带样式的空白 docx"""

    # 先创建空白文档
    doc = Document()

    # 保存到临时文件
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    doc.save(tmp_path)

    # 读取提取的样式
    with open(styles_xml_path, "rb") as f:
        new_styles = etree.parse(f).getroot()

    # 解压 docx，修改 styles.xml，重新打包
    with tempfile.TemporaryDirectory() as tmpdir:
        # 解压
        with zipfile.ZipFile(tmp_path, "r") as zf:
            zf.extractall(tmpdir)

        # 读取原 styles.xml
        orig_styles_path = os.path.join(tmpdir, "word", "styles.xml")
        with open(orig_styles_path, "rb") as f:
            orig_root = etree.parse(f).getroot()

        # 合并样式：把新样式添加到原样式中
        # 先删除同名样式
        existing_ids = set()
        for style in new_styles.findall(".//w:style", NSMAP):
            style_id = style.get(f"{{{NSMAP['w']}}}styleId")
            existing_ids.add(style_id)

        for style in orig_root.findall(".//w:style", NSMAP):
            style_id = style.get(f"{{{NSMAP['w']}}}styleId")
            if style_id in existing_ids:
                orig_root.remove(style)

        # 添加新样式
        for style in new_styles.findall(".//w:style", NSMAP):
            orig_root.append(deepcopy(style))

        # 更新 docDefaults
        new_defaults = new_styles.find(".//w:docDefaults", NSMAP)
        if new_defaults is not None:
            old_defaults = orig_root.find(".//w:docDefaults", NSMAP)
            if old_defaults is not None:
                orig_root.remove(old_defaults)
            orig_root.insert(0, deepcopy(new_defaults))

        # 写回 styles.xml
        tree = etree.ElementTree(orig_root)
        tree.write(orig_styles_path, encoding="UTF-8", xml_declaration=True)

        # 重新打包
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for root_dir, dirs, files in os.walk(tmpdir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, tmpdir)
                    zf.write(file_path, arcname)

    # 清理临时文件
    os.unlink(tmp_path)

    return output_path


def set_table_border(table, border_color="000000", border_size=4):
    """给表格设置边框

    Args:
        table: python-docx Table 对象
        border_color: 边框颜色（十六进制，默认黑色）
        border_size: 边框粗细（单位：1/8磅，4=0.5磅）
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # 创建边框元素
    tblBorders = OxmlElement("w:tblBorders")

    # 六种边框：top, left, bottom, right, insideH, insideV
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(border_size))
        border.set(qn("w:color"), border_color)
        border.set(qn("w:space"), "0")
        tblBorders.append(border)

    # 移除旧边框设置
    old_borders = tblPr.find(qn("w:tblBorders"))
    if old_borders is not None:
        tblPr.remove(old_borders)

    tblPr.append(tblBorders)


def parse_table_row(line):
    """解析表格行，返回单元格列表"""
    # 去掉首尾的 |
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    # 分割单元格
    cells = [cell.strip() for cell in line.split("|")]
    return cells


def is_separator_row(line):
    """检查是否是表格分隔行 |---|---|"""
    line = line.strip()
    if not line.startswith("|"):
        return False
    # 去掉 | 后检查是否只有 - : 空格
    content = line.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
    return len(content) == 0


def clean_markdown_text(text):
    """清理 Markdown 格式"""
    # 去除 markdown 格式
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # **bold**
    text = re.sub(r"\*(.+?)\*", r"\1", text)  # *italic*
    text = re.sub(r"`(.+?)`", r"\1", text)  # `code`
    text = re.sub(r"\$(.+?)\$", r"\1", text)  # $math$ (简单处理)
    return text


def parse_list_item(line):
    """解析列表项，返回 (缩进级别, 内容)"""
    # 计算缩进（空格或 tab）
    stripped = line.lstrip()
    indent = len(line) - len(stripped)
    indent_level = indent // 2  # 每 2 空格一级

    # 去掉列表标记
    if (
        stripped.startswith("- ")
        or stripped.startswith("* ")
        and not stripped.startswith("**")
        or stripped.startswith("> ")
    ):
        content = stripped[2:]
    elif re.match(r"^\d+\.\s", stripped):
        content = re.sub(r"^\d+\.\s", "", stripped)
    else:
        content = stripped

    return indent_level, content.strip()


def merge_list_items(items):
    """合并列表项为段落文本

    智能处理：如果前一项以冒号结尾，则不加分号分隔
    """
    if not items:
        return ""
    if len(items) == 1:
        return items[0]

    result = items[0]
    for item in items[1:]:
        # 如果前面以冒号结尾，直接连接
        if result.endswith("：") or result.endswith(":"):
            result += item
        else:
            result += "；" + item

    return result


def parse_markdown(md_content):
    """解析 Markdown，返回元素列表"""
    elements = []
    lines = md_content.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行跳过
        if not stripped:
            i += 1
            continue

        # 分隔线 ---
        if stripped == "---":
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            elements.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # 表名：以 "表" + 数字 开头
        if re.match(r"^表\d+", stripped):
            elements.append({"type": "table_title", "text": stripped})
            i += 1
            continue

        # 图名：以 "图" + 数字 开头
        if re.match(r"^图\d+", stripped):
            elements.append({"type": "figure_title", "text": stripped})
            i += 1
            continue

        # 表格：以 | 开头
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # 解析表格
            if len(table_lines) >= 2:
                headers = []
                rows = []

                # 第一行是表头
                headers = parse_table_row(table_lines[0])

                # 第二行是分隔符，跳过
                # 后面是数据行
                for tl in table_lines[2:]:
                    if not is_separator_row(tl):
                        rows.append(parse_table_row(tl))

                elements.append({"type": "table", "headers": headers, "rows": rows})
            continue

        # 列表项：以 - 、* 、> 或数字. 开头
        # 注意：区分 "* 列表项"（单星号+空格）和 "**加粗**"（双星号，不是列表）
        is_list_start = (
            stripped.startswith("- ")
            or (stripped.startswith("* ") and not stripped.startswith("**"))
            or stripped.startswith("> ")
            or re.match(r"^\d+\.\s", stripped)
        )

        if is_list_start:
            # 收集连续的列表项
            list_items = []
            while i < len(lines):
                current = lines[i]
                current_stripped = current.strip()

                # 空行结束列表
                if not current_stripped:
                    break

                # 检查是否是列表项（包括缩进的子项）
                is_list_item = (
                    current_stripped.startswith("- ")
                    or (current_stripped.startswith("* ") and not current_stripped.startswith("**"))
                    or current_stripped.startswith("> ")
                    or re.match(r"^\d+\.\s", current_stripped)
                    or (current.startswith("  ") and list_items)  # 缩进的续行
                )

                if not is_list_item:
                    break

                indent_level, content = parse_list_item(current)
                if content:  # 只添加非空内容
                    list_items.append((indent_level, content))
                i += 1

            # 把列表项转换为段落
            # 策略：每个列表项作为一个段落，缩进项合并到上一项
            if list_items:
                current_para = []
                for indent_level, content in list_items:
                    clean_content = clean_markdown_text(content)
                    if indent_level == 0:
                        # 新的一级列表项
                        if current_para:
                            # 保存之前的段落
                            elements.append({"type": "paragraph", "text": merge_list_items(current_para)})
                        current_para = [clean_content]
                    else:
                        # 缩进的子项，合并到当前段落
                        current_para.append(clean_content)

                # 保存最后一个段落
                if current_para:
                    elements.append({"type": "paragraph", "text": merge_list_items(current_para)})
            continue

        # 普通段落
        para_lines = [line]
        i += 1
        # 继续收集段落内容，直到遇到特殊行
        while i < len(lines):
            next_line = lines[i]
            next_stripped = next_line.strip()

            # 空行结束段落
            if not next_stripped:
                break

            # 特殊行结束段落
            if (
                re.match(r"^#{1,4}\s", next_line)  # 标题
                or next_stripped.startswith("|")  # 表格
                or next_stripped.startswith("- ")  # 列表
                or (next_stripped.startswith("* ") and not next_stripped.startswith("**"))
                or next_stripped.startswith("> ")
                or re.match(r"^\d+\.\s", next_stripped)
                or next_stripped == "---"
            ):
                break

            para_lines.append(next_line)
            i += 1

        text = " ".join(para_lines)
        text = clean_markdown_text(text)

        if text.strip():  # 只添加非空段落
            elements.append({"type": "paragraph", "text": text})

    return elements


def convert_md_to_docx(md_path, styles_xml_path, output_path, config=None):
    """转换 Markdown 到 Docx"""

    print(f"📖 读取: {md_path}")

    # 读取 markdown
    with open(md_path, encoding="utf-8") as f:
        md_content = f.read()

    # 解析
    elements = parse_markdown(md_content)
    print(f"📊 解析: {len(elements)} 个元素")

    # 创建带样式的 docx
    create_docx_with_styles(styles_xml_path, output_path)

    # 打开并写入内容
    doc = Document(output_path)

    # 确定正文样式名
    body_style = "ZDWP正文"
    if config and config.get("body_style_name"):
        body_style = config["body_style_name"]

    # 表格内容样式
    table_cell_style = "ZDWP表格内容"
    # 表名样式
    table_title_style = "ZDWP表名"
    # 图名样式
    figure_title_style = "ZDWP图名"

    # 检查样式是否存在
    available_styles = {s.name for s in doc.styles}
    if body_style not in available_styles:
        body_style = "Normal"
        print("⚠️  使用 Normal 作为正文样式")

    if table_cell_style not in available_styles:
        table_cell_style = body_style
        print(f"⚠️  使用 {body_style} 作为表格内容样式")

    if table_title_style not in available_styles:
        table_title_style = body_style
        print(f"⚠️  使用 {body_style} 作为表名样式")

    if figure_title_style not in available_styles:
        figure_title_style = body_style
        print(f"⚠️  使用 {body_style} 作为图名样式")

    # 样式映射
    heading_styles = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
    }

    # 写入内容
    for elem in elements:
        if elem["type"] == "heading":
            level = elem["level"]
            style_name = heading_styles.get(level, "Heading 4")
            try:
                doc.add_paragraph(elem["text"], style=style_name)
            except KeyError:
                doc.add_heading(elem["text"], level=level)

        elif elem["type"] == "paragraph":
            try:
                doc.add_paragraph(elem["text"], style=body_style)
            except KeyError:
                doc.add_paragraph(elem["text"])

        elif elem["type"] == "table_title":
            try:
                doc.add_paragraph(elem["text"], style=table_title_style)
            except KeyError:
                doc.add_paragraph(elem["text"])

        elif elem["type"] == "figure_title":
            try:
                doc.add_paragraph(elem["text"], style=figure_title_style)
            except KeyError:
                doc.add_paragraph(elem["text"])

        elif elem["type"] == "table":
            headers = elem["headers"]
            rows = elem["rows"]

            # 计算行列数
            num_cols = len(headers)
            num_rows = 1 + len(rows)  # 表头 + 数据行

            # 创建表格
            table = doc.add_table(rows=num_rows, cols=num_cols)

            # 设置表格边框
            set_table_border(table)

            # 填充表头
            header_row = table.rows[0]
            for j, header in enumerate(headers):
                cell = header_row.cells[j]
                # 垂直居中
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # 清空默认段落，设置样式
                cell.text = ""
                para = cell.paragraphs[0]
                para.text = header
                # 水平居中
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    para.style = table_cell_style
                except KeyError:
                    pass

            # 填充数据行
            for i, row_data in enumerate(rows):
                row = table.rows[i + 1]
                for j, cell_text in enumerate(row_data):
                    if j < num_cols:
                        cell = row.cells[j]
                        # 垂直居中
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.text = ""
                        para = cell.paragraphs[0]
                        para.text = cell_text
                        # 水平居中
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        try:
                            para.style = table_cell_style
                        except KeyError:
                            pass

    # 保存
    doc.save(output_path)
    print(f"✅ 输出: {output_path}")

    return output_path


def get_finder_selection():
    """获取 Finder 选中的文件"""
    import subprocess

    script = """
    tell application "Finder"
        set sel to selection
        if (count of sel) > 0 then
            return POSIX path of (item 1 of sel as alias)
        end if
    end tell
    """
    result = subprocess.run(["osascript", "-e", script], capture_output=True, text=True)
    return result.stdout.strip()


# 默认模板路径
DEFAULT_TEMPLATE = "/Users/tianli/Downloads/归档/其他文档/template.docx"


def main():
    parser = argparse.ArgumentParser(
        description="Markdown 转 Docx（样式复刻版）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 提取样式
  python md_docx_heading_template.py extract 模板.docx

  # 转换（Raycast 调用：选中 .md 文件即可）
  python md_docx_heading_template.py input.md

  # 指定模板
  python md_docx_heading_template.py input.md -t 模板.docx -o output.docx
        """,
    )

    parser.add_argument("input", nargs="?", help="Markdown 文件或命令 (extract/convert)")
    parser.add_argument("template_or_md", nargs="?", help="模板 docx (extract) 或 md 文件 (convert)")
    parser.add_argument("-t", "--template", help="模板 docx 文件")
    parser.add_argument("-s", "--styles", help="样式 XML 文件")
    parser.add_argument("-o", "--output", help="输出文件")

    args = parser.parse_args()

    # 无参数时从 Finder 获取选中的 .md 文件
    if not args.input:
        finder_file = get_finder_selection()
        if finder_file and finder_file.endswith(".md"):
            args.input = finder_file
            print(f"📄 从 Finder 获取: {os.path.basename(finder_file)}")
        else:
            print("❌ 请在 Finder 中选择一个 .md 文件")
            sys.exit(1)

    # 命令: extract
    if args.input == "extract":
        if not args.template_or_md:
            print("❌ 用法: python md_docx_heading_template.py extract 模板.docx")
            sys.exit(1)
        extract_styles_xml(args.template_or_md)
        return

    # 命令: convert
    if args.input == "convert":
        if not args.template_or_md:
            print("❌ 用法: python md_docx_heading_template.py convert input.md -o output.docx")
            sys.exit(1)

        md_path = args.template_or_md
        styles_path = args.styles or os.path.join(DEFAULT_STYLES_DIR, "heading_styles.xml")
        output_path = args.output or os.path.splitext(md_path)[0] + ".docx"

        if not os.path.exists(styles_path):
            print(f"❌ 样式文件不存在: {styles_path}")
            print("   请先运行: python md_docx_heading_template.py extract 模板.docx")
            sys.exit(1)

        # 读取配置
        config_path = os.path.join(os.path.dirname(styles_path), "styles_config.json")
        config = None
        if os.path.exists(config_path):
            with open(config_path, encoding="utf-8") as f:
                config = json.load(f)

        convert_md_to_docx(md_path, styles_path, output_path, config)
        return

    # 一步完成: input.md -t 模板.docx -o output.docx
    if args.input.endswith(".md"):
        md_path = args.input

        # 使用指定模板或默认模板
        template = args.template or DEFAULT_TEMPLATE

        if os.path.exists(template):
            # 先提取样式
            with tempfile.TemporaryDirectory() as tmpdir:
                print(f"📋 使用模板: {os.path.basename(template)}")
                styles_path, config = extract_styles_xml(template, tmpdir)
                output_path = args.output or os.path.splitext(md_path)[0] + ".docx"
                convert_md_to_docx(md_path, styles_path, output_path, config)
        else:
            # 使用已提取的样式
            styles_path = args.styles or os.path.join(DEFAULT_STYLES_DIR, "heading_styles.xml")
            if not os.path.exists(styles_path):
                print(f"❌ 样式文件不存在: {styles_path}")
                print(f"❌ 默认模板不存在: {DEFAULT_TEMPLATE}")
                print("   请指定模板: -t 模板.docx")
                sys.exit(1)

            config_path = os.path.join(os.path.dirname(styles_path), "styles_config.json")
            config = None
            if os.path.exists(config_path):
                with open(config_path, encoding="utf-8") as f:
                    config = json.load(f)

            output_path = args.output or os.path.splitext(md_path)[0] + ".docx"
            convert_md_to_docx(md_path, styles_path, output_path, config)
        return

    parser.print_help()


if __name__ == "__main__":
    main()
