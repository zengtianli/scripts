#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title xlsx-lowercase
# @raycast.mode fullOutput
# @raycast.icon 📊
# @raycast.packageName Scripts
# @raycast.description Lowercase Excel data
# -*- coding: utf-8 -*-
"""
Office文档文本小写化工具

功能：将 Office 文档（Word、Excel）中的英文字母转为小写

支持格式：
- Word文档：.docx
- Excel文档：.xlsx, .xlsm

处理规则：
- Word：转换所有段落和表格中的文本
- Excel：第1行说明、第2行表头不处理，第3行起数据转小写
- 生成新文件（原文件名_lower.xxx）
"""

import sys
from pathlib import Path

# 添加 _lib 到搜索路径
sys.path.insert(0, str(Path(__file__).parent / "_lib"))

from finder import get_input_files


def process_word(input_path, output_path):
    """
    处理Word文档，将文本转为小写
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        
    Returns:
        bool: 处理是否成功
    """
    try:
        from docx import Document
    except ImportError:
        print("❌ 错误：缺少python-docx库")
        print("请安装：pip install python-docx")
        return False
    
    try:
        print(f"📖 正在读取Word文档: {input_path.name}")
        doc = Document(input_path)
        
        # 统计信息
        paragraphs_changed = 0
        table_cells_changed = 0
        
        print(f"🔄 正在处理段落...")
        # 处理段落
        for para in doc.paragraphs:
            if para.text:
                original_text = para.text
                # 处理段落中的每个run以保持格式
                for run in para.runs:
                    if run.text:
                        old_text = run.text
                        new_text = old_text.lower()
                        if old_text != new_text:
                            run.text = new_text
                            paragraphs_changed += 1
        
        print(f"🔄 正在处理表格...")
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                old_text = run.text
                                new_text = old_text.lower()
                                if old_text != new_text:
                                    run.text = new_text
                                    table_cells_changed += 1
        
        # 保存文档
        print(f"💾 正在保存文件...")
        doc.save(output_path)
        
        print(f"✅ 处理完成！")
        print(f"   - 修改段落: {paragraphs_changed} 个")
        print(f"   - 修改表格单元格: {table_cells_changed} 个")
        print(f"   - 输出文件: {output_path.name}")
        
        return True
        
    except Exception as e:
        print(f"❌ 处理Word文档失败：{e}")
        import traceback
        traceback.print_exc()
        return False


def process_excel(input_path, output_path):
    """
    处理Excel文件，将数据行的英文转为小写
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径
        
    Returns:
        bool: 处理是否成功
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        print("❌ 错误：缺少openpyxl库")
        print("请安装：pip install openpyxl")
        return False
    
    try:
        print(f"📖 正在读取Excel文件: {input_path.name}")
        wb = load_workbook(input_path)
        
        # 统计信息
        stats = {
            'total_sheets': len(wb.sheetnames),
            'processed_sheets': 0,
            'total_cells_processed': 0,
            'total_cells_changed': 0
        }
        
        print(f"🔄 正在处理Excel文件...")
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            # 跳过空sheet或数据不足的sheet
            if ws.max_row < 3:
                print(f"  ⏭️  跳过 '{sheet_name}': 数据行数不足")
                continue
            
            cells_processed = 0
            cells_changed = 0
            
            # 获取最大行列数
            max_row = ws.max_row
            max_col = ws.max_column
            
            # 如果行数太多，显示进度
            show_progress = max_row > 1000
            
            # 从第3行开始处理（第1行说明，第2行表头）
            for row_idx in range(3, max_row + 1):
                for col_idx in range(1, max_col + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    old_value = cell.value
                    
                    if old_value is not None:
                        # 如果是字符串，转小写
                        if isinstance(old_value, str):
                            new_value = old_value.lower()
                            
                            # 如果值发生了变化
                            if old_value != new_value:
                                cell.value = new_value
                                cells_changed += 1
                        
                        cells_processed += 1
                
                # 显示进度
                if show_progress and (row_idx - 2) % 1000 == 0:
                    print(f"    进度: {row_idx - 2}/{max_row - 2} 行...", end='\r')
            
            if cells_processed > 0:
                print(f"  ✓ '{sheet_name}': 处理 {cells_processed} 个单元格, 修改 {cells_changed} 个")
                stats['processed_sheets'] += 1
                stats['total_cells_processed'] += cells_processed
                stats['total_cells_changed'] += cells_changed
        
        # 保存到新文件
        print(f"💾 正在保存文件...")
        wb.save(output_path)
        
        # 输出统计信息
        print(f"✅ 处理完成！")
        print(f"   - 共处理 {stats['processed_sheets']}/{stats['total_sheets']} 个工作表")
        print(f"   - 共处理 {stats['total_cells_processed']} 个单元格")
        print(f"   - 共修改 {stats['total_cells_changed']} 个单元格")
        print(f"   - 输出文件: {output_path.name}")
        
        return True
        
    except Exception as e:
        print(f"❌ 处理Excel文件失败：{e}")
        import traceback
        traceback.print_exc()
        return False


def process_file(input_file):
    """
    根据文件类型处理Office文档
    
    Args:
        input_file: 文件路径
        
    Returns:
        bool: 处理是否成功
    """
    input_path = Path(input_file)
    
    # 检查文件是否存在
    if not input_path.exists():
        print(f"❌ 错误：文件不存在 - {input_file}")
        return False
    
    # 获取文件扩展名
    suffix = input_path.suffix.lower()
    
    # 生成输出文件名
    output_path = input_path.parent / f"{input_path.stem}_lower{input_path.suffix}"
    
    # 根据文件类型调用不同的处理函数
    if suffix == '.docx':
        return process_word(input_path, output_path)
    elif suffix in ['.xlsx', '.xlsm']:
        return process_excel(input_path, output_path)
    else:
        print(f"❌ 错误：不支持的文件格式 '{suffix}'")
        print(f"支持的格式：.docx, .xlsx, .xlsm")
        return False


if __name__ == "__main__":
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext=['docx', 'xlsx', 'xlsm'], allow_multiple=False)
    
    if not files:
        print("❌ 错误：缺少文件名参数")
        print("\n使用方法：")
        print("    python3 office_lowercase.py 文件名.docx")
        print("    python3 office_lowercase.py 文件名.xlsx")
        print("    或在 Finder 中选择文件后运行")
        print("\n支持格式：")
        print("    - Word文档：.docx")
        print("    - Excel文档：.xlsx, .xlsm")
        sys.exit(1)
    
    input_file = files[0]
    
    print("=" * 50)
    print("Office文档文本小写化工具")
    print("=" * 50)
    
    success = process_file(input_file)
    
    if success:
        print("\n🎉 全部完成！")
    else:
        print("\n❌ 处理失败，请检查错误信息")
        sys.exit(1)
