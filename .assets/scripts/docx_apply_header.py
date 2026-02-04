#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-header
# @raycast.mode fullOutput
# @raycast.icon 📄
# @raycast.packageName Scripts
# @raycast.description Apply header to Word
"""
页眉设置工具
为Word文档添加标准页眉：文档标题（左对齐）+ 章节引用（右对齐）

用法:
    python3 apply_header.py <input.docx> [文档标题]
    
示例:
    python3 apply_header.py document.docx
    python3 apply_header.py document.docx "浙江省水利工程报告"
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from common_utils import get_input_files

from common_utils import get_input_files

from common_utils import get_input_files


def create_element(name):
    """创建XML元素"""
    return OxmlElement(name)


def create_attribute(element, name, value):
    """设置XML属性"""
    element.set(qn(name), value)


def get_first_heading1_text(doc):
    """
    获取文档中第一个标题1的文本
    
    Args:
        doc: Document对象
        
    Returns:
        str: 标题文本，如果没有找到返回空字符串
    """
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name
            if 'Heading 1' in style_name or '标题 1' in style_name or style_name == '标题1':
                text = paragraph.text.strip()
                if text:
                    return text
    return ""


def add_styleref_field(paragraph, style_name="标题 1"):
    """
    在段落中添加STYLEREF域（引用指定样式的文本）
    
    Args:
        paragraph: Paragraph对象
        style_name: 样式名称（默认"标题 1"）
    """
    run = paragraph.add_run()
    
    # 创建 fldChar begin
    fldChar_begin = create_element('w:fldChar')
    create_attribute(fldChar_begin, 'w:fldCharType', 'begin')
    
    # 创建 instrText (STYLEREF)
    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' STYLEREF "{style_name}" \\* MERGEFORMAT '
    
    # 创建 fldChar separate
    fldChar_separate = create_element('w:fldChar')
    create_attribute(fldChar_separate, 'w:fldCharType', 'separate')
    
    # 创建占位符文本
    placeholder = create_element('w:t')
    placeholder.text = "章节标题"
    
    # 创建 fldChar end
    fldChar_end = create_element('w:fldChar')
    create_attribute(fldChar_end, 'w:fldCharType', 'end')
    
    # 组装域代码
    run._r.append(fldChar_begin)
    
    run2 = paragraph.add_run()
    run2._r.append(instrText)
    
    run3 = paragraph.add_run()
    run3._r.append(fldChar_separate)
    
    run4 = paragraph.add_run()
    run4._r.append(placeholder)
    
    run5 = paragraph.add_run()
    run5._r.append(fldChar_end)


def set_header(doc, doc_title=""):
    """
    设置页眉：文档标题（左）+ 章节引用（右）
    
    Args:
        doc: Document对象
        doc_title: 文档标题，如果为空则自动获取第一个标题1
    """
    # 如果没有提供标题，尝试获取第一个标题1
    if not doc_title:
        doc_title = get_first_heading1_text(doc)
    
    if not doc_title:
        doc_title = "文档标题"
    
    # 为所有节设置页眉
    for i, section in enumerate(doc.sections):
        header = section.header
        header.is_linked_to_previous = False
        
        # 清空现有内容
        for paragraph in header.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # 创建页眉段落
        para = header.add_paragraph()
        
        # 设置段落制表位
        pPr = para._element.get_or_add_pPr()
        tabs = create_element('w:tabs')
        
        # 右对齐制表位
        tab_right = create_element('w:tab')
        create_attribute(tab_right, 'w:val', 'right')
        create_attribute(tab_right, 'w:pos', '9360')  # 约16.5cm
        tabs.append(tab_right)
        
        pPr.append(tabs)
        
        # 添加文档标题（左对齐）
        run_title = para.add_run(doc_title)
        run_title.font.size = Pt(9)
        run_title.font.name = '宋体'
        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加制表符
        para.add_run('\t')
        
        # 添加章节引用域（右对齐）
        add_styleref_field(para, "标题 1")
        
        # 添加页眉下划线
        pBdr = create_element('w:pBdr')
        bottom = create_element('w:bottom')
        create_attribute(bottom, 'w:val', 'single')
        create_attribute(bottom, 'w:sz', '4')
        create_attribute(bottom, 'w:space', '1')
        create_attribute(bottom, 'w:color', '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    return doc_title


def apply_header(input_file, doc_title=""):
    """
    应用页眉设置
    
    Args:
        input_file: 输入文件路径
        doc_title: 文档标题（可选）
    """
    input_path = Path(input_file)
    
    # 检查文件是否存在
    if not input_path.exists():
        print(f"❌ 错误: 文件不存在: {input_file}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"❌ 错误: 只支持 .docx 文件")
        sys.exit(1)
    
    print(f"🔄 正在处理文件: {input_path.name}")
    
    # 加载文档
    try:
        doc = Document(str(input_path))
    except Exception as e:
        print(f"❌ 错误: 无法打开文件: {e}")
        sys.exit(1)
    
    # 备份原文件
    backup_path = input_path.with_suffix('.docx.backup')
    try:
        import shutil
        shutil.copy2(str(input_path), str(backup_path))
        print(f"ℹ️ 已备份原文件: {backup_path.name}")
    except Exception as e:
        print(f"⚠️ 备份失败: {e}")
    
    # 设置页眉
    print(f"🔄 正在设置页眉...")
    try:
        if not doc_title:
            doc_title = get_first_heading1_text(doc)
            if doc_title:
                print(f'ℹ️ 自动获取文档标题: {doc_title[:30]}...')
            else:
                doc_title = "文档标题"
                print(f'ℹ️ 未找到标题1，使用默认标题')
        
        actual_title = set_header(doc, doc_title)
        print(f"✅ 页眉设置完成!")
        print(f"   - 左侧: {actual_title[:20]}...")
        print(f"   - 右侧: 章节引用（自动跟随标题1）")
    except Exception as e:
        print(f"❌ 页眉设置失败: {e}")
        sys.exit(1)
    
    # 保存文档
    try:
        doc.save(str(input_path))
        print(f"✅ 已保存: {input_path.name}")
        if backup_path.exists():
            print(f"ℹ️ 如需恢复，请使用备份文件: {backup_path.name}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")
        sys.exit(1)


def main():
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext='docx', allow_multiple=False)
    
    if not files:
        print("页眉设置工具")
        print("\n用法: python3 apply_header.py <input.docx> [文档标题]")
        print("      或在 Finder 中选择 .docx 文件后运行")
        print("\n示例:")
        print("  python3 apply_header.py document.docx")
        print('  python3 apply_header.py document.docx "浙江省水利工程报告"')
        print("\n功能:")
        print("  - 左侧: 文档标题（自动获取第一个标题1，或手动指定）")
        print("  - 右侧: 章节引用（自动跟随当前页的标题1）")
        print("  - 带下划线")
        sys.exit(1)
    
    input_file = files[0]
    doc_title = sys.argv[2] if len(sys.argv) > 2 else ""
    
    apply_header(input_file, doc_title)


if __name__ == "__main__":
    main()

