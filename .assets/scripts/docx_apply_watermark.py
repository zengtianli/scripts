#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-watermark
# @raycast.mode fullOutput
# @raycast.icon 📄
# @raycast.packageName Scripts
# @raycast.description Apply watermark
"""
水印设置工具
为Word文档添加对角线文字水印

用法:
    python3 apply_watermark.py <input.docx> [水印文字] [字体] [字号]
    
示例:
    python3 apply_watermark.py document.docx
    python3 apply_watermark.py document.docx "ZDWP"
    python3 apply_watermark.py document.docx "ZDWP" "Arial" 96
"""

import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from common_utils import get_input_files

from common_utils import get_input_files

# 默认水印设置
DEFAULT_WATERMARK_TEXT = "ZDWP"
DEFAULT_WATERMARK_FONT = "Arial"
DEFAULT_WATERMARK_SIZE = 96


def add_watermark(doc, text, font, size):
    """
    添加对角线文字水印
    使用完整的 VML XML 结构
    
    Args:
        doc: Document对象
        text: 水印文字
        font: 字体名称
        size: 字号
    """
    from lxml import etree
    
    # 命名空间
    ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    ns_v = '{urn:schemas-microsoft-com:vml}'
    ns_o = '{urn:schemas-microsoft-com:office:office}'
    ns_r = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
    
    # 为所有节添加水印
    for idx, section in enumerate(doc.sections):
        header = section.header
        header.is_linked_to_previous = False
        
        # 确保有段落
        if not header.paragraphs:
            header.add_paragraph()
        
        # 获取 header 的 XML
        hdr = header._element
        
        # 查找第一个段落
        p_elements = hdr.findall(f'.//{ns_w}p')
        if not p_elements:
            continue
        
        p = p_elements[0]
        
        # 创建 r 元素
        r = etree.SubElement(p, f'{ns_w}r')
        
        # 创建 pict 元素
        pict = etree.SubElement(r, f'{ns_w}pict')
        
        # 首先创建 shapetype（Word 需要这个定义）
        shapetype = etree.SubElement(pict, f'{ns_v}shapetype')
        shapetype.set('id', '_x0000_t136')
        shapetype.set('coordsize', '21600,21600')
        shapetype.set(f'{ns_o}spt', '136')
        shapetype.set('adj', '10800')
        shapetype.set('path', 'm@7,l@8,m@5,21600l@6,21600e')
        
        # 添加 formulas
        formulas = etree.SubElement(shapetype, f'{ns_v}formulas')
        formula_eqns = [
            'sum #0 0 10800', 'prod #0 2 1', 'sum 21600 0 @1', 'sum 0 0 @2',
            'sum 21600 0 @3', 'if @0 @3 0', 'if @0 21600 @1', 'if @0 0 @2',
            'if @0 @4 21600', 'mid @5 @6', 'mid @8 @5', 'mid @7 @8',
            'mid @6 @7', 'sum @6 0 @5'
        ]
        for eqn in formula_eqns:
            f = etree.SubElement(formulas, f'{ns_v}f')
            f.set('eqn', eqn)
        
        # 添加 path
        path = etree.SubElement(shapetype, f'{ns_v}path')
        path.set('textpathok', 't')
        path.set(f'{ns_o}connecttype', 'custom')
        path.set(f'{ns_o}connectlocs', '@9,0;@10,10800;@11,21600;@12,10800')
        path.set(f'{ns_o}connectangles', '270,180,90,0')
        
        # 添加 textpath 到 shapetype
        tp = etree.SubElement(shapetype, f'{ns_v}textpath')
        tp.set('on', 't')
        tp.set('fitshape', 't')
        
        # 添加 handles
        handles = etree.SubElement(shapetype, f'{ns_v}handles')
        h = etree.SubElement(handles, f'{ns_v}h')
        h.set('position', '#0,bottomRight')
        h.set('xrange', '6629,14971')
        
        # 添加 lock
        lock = etree.SubElement(shapetype, f'{ns_o}lock')
        lock.set(f'{ns_v}ext', 'edit')
        lock.set('text', 't')
        lock.set('shapetype', 't')
        
        # 创建实际的水印 shape
        shape = etree.SubElement(pict, f'{ns_v}shape')
        shape.set('id', f'PowerPlusWaterMarkObject{idx}')
        shape.set(f'{ns_o}spid', f'_x0000_s{2049 + idx}')
        shape.set('type', '#_x0000_t136')
        shape.set('style', 
                  'position:absolute;'
                  'margin-left:0;'
                  'margin-top:0;'
                  'width:527.85pt;'
                  'height:131.95pt;'
                  'z-index:-251656192;'
                  'mso-wrap-edited:f;'
                  'mso-position-horizontal:center;'
                  'mso-position-horizontal-relative:margin;'
                  'mso-position-vertical:center;'
                  'mso-position-vertical-relative:margin;'
                  'rotation:315')
        shape.set(f'{ns_o}allowincell', 'f')
        shape.set('fillcolor', 'silver')
        shape.set('stroked', 'f')
        
        # 添加 fill 元素（设置透明度）
        fill = etree.SubElement(shape, f'{ns_v}fill')
        fill.set('opacity', '.5')
        
        # 添加 textpath 元素（设置文字）
        textpath = etree.SubElement(shape, f'{ns_v}textpath')
        textpath.set('style', f'font-family:"{font}";font-size:{size}pt')
        textpath.set('string', text)


def apply_watermark(input_file, text=DEFAULT_WATERMARK_TEXT, 
                    font=DEFAULT_WATERMARK_FONT, size=DEFAULT_WATERMARK_SIZE):
    """
    应用水印设置
    
    Args:
        input_file: 输入文件路径
        text: 水印文字
        font: 字体名称
        size: 字号
    """
    input_path = Path(input_file)
    
    # 检查文件是否存在
    if not input_path.exists():
        print(f"❌ 错误: 文件不存在: {input_file}")
        sys.exit(1)
    
    if input_path.suffix.lower() != '.docx':
        print(f"❌ 错误: 只支持 .docx 文件")
        sys.exit(1)
    
    print(f"🔄 正在处理文件: {input_path.name}")
    
    # 加载文档
    try:
        doc = Document(str(input_path))
    except Exception as e:
        print(f"❌ 错误: 无法打开文件: {e}")
        sys.exit(1)
    
    # 备份原文件
    backup_path = input_path.with_suffix('.docx.backup')
    try:
        import shutil
        shutil.copy2(str(input_path), str(backup_path))
        print(f"ℹ️ 已备份原文件: {backup_path.name}")
    except Exception as e:
        print(f"⚠️ 备份失败: {e}")
    
    # 添加水印
    print(f"🔄 正在添加水印...")
    try:
        add_watermark(doc, text, font, size)
        print(f"✅ 水印添加完成!")
        print(f"   - 文字: {text}")
        print(f"   - 字体: {font}")
        print(f"   - 字号: {size}pt")
        print(f"   - 方向: 对角线（-45°）")
    except Exception as e:
        print(f"❌ 水印添加失败: {e}")
        sys.exit(1)
    
    # 保存文档
    try:
        doc.save(str(input_path))
        print(f"✅ 已保存: {input_path.name}")
        if backup_path.exists():
            print(f"ℹ️ 如需恢复，请使用备份文件: {backup_path.name}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")
        sys.exit(1)


def main():
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext='docx', allow_multiple=False)
    
    if not files:
        print("水印设置工具")
        print("\n用法: python3 apply_watermark.py <input.docx> [水印文字] [字体] [字号]")
        print("      或在 Finder 中选择 .docx 文件后运行")
        print("\n示例:")
        print("  python3 apply_watermark.py document.docx")
        print('  python3 apply_watermark.py document.docx "ZDWP"')
        print('  python3 apply_watermark.py document.docx "ZDWP" "Arial" 96')
        print("\n默认设置:")
        print(f"  - 文字: {DEFAULT_WATERMARK_TEXT}")
        print(f"  - 字体: {DEFAULT_WATERMARK_FONT}")
        print(f"  - 字号: {DEFAULT_WATERMARK_SIZE}pt")
        print("  - 方向: 对角线（-45°）")
        print("  - 颜色: 银灰色，30%透明度")
        sys.exit(1)
    
    input_file = files[0]
    text = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_WATERMARK_TEXT
    font = sys.argv[3] if len(sys.argv) > 3 else DEFAULT_WATERMARK_FONT
    size = int(sys.argv[4]) if len(sys.argv) > 4 else DEFAULT_WATERMARK_SIZE
    
    apply_watermark(input_file, text, font, size)


if __name__ == "__main__":
    main()

