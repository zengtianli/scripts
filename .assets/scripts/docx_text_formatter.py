#!/usr/bin/env python3
# @raycast.schemaVersion 1
# @raycast.title docx-text-format
# @raycast.mode fullOutput
# @raycast.icon 📄
# @raycast.packageName Scripts
# @raycast.description Format text in Word
# @raycast.argument1 { "type": "text", "placeholder": "文件路径", "optional": true }
# -*- coding: utf-8 -*-
"""
文本格式自动修复工具 - DOCX版本
支持处理Word文档中的格式问题

功能列表：
1. 双引号格式修复
   将所有双引号统一为中文标准引号："和"
   奇数个引号 → " (中文左双引号 U+201C)
   偶数个引号 → " (中文右双引号 U+201D)

2. 英文标点符号转中文
   , → ，  (逗号)
   : → ：  (冒号)
   ; → ；  (分号)
   ! → ！  (感叹号)
   ? → ？  (问号)
   ( → （  (左括号)
   ) → ）  (右括号)

3. 中文单位转标准符号
   面积：平方米 → m²、平方公里 → km²
   体积：立方米 → m³、立方厘米 → cm³
   长度：公里 → km、厘米 → cm、毫米 → mm
   质量：公斤 → kg、毫克 → mg
   容量：毫升 → mL、微升 → μL
   时间：小时 → h、分钟 → min、秒钟 → s
   温度：摄氏度 → ℃、华氏度 → ℉

4. 单位上标格式化
   m2 → m²、m3 → m³、km2 → km²、km3 → km³

使用方法：
    python3 text_formatter_docx.py 文件名.docx
"""

import re
import sys
from pathlib import Path
from docx import Document

# ===== 配置选项 =====
# 注释掉下面这行可以启用页脚处理
SKIP_FOOTER = True
# SKIP_FOOTER = False  # 取消注释这行可以处理页脚


def fix_quotes(content):
    """
    替换所有双引号为中文标准引号
    奇数位置 → " (左引号)
    偶数位置 → " (右引号)
    """
    # 匹配所有可能的双引号类型
    # ", ", ", ", 「, 」以及英文双引号"
    quote_pattern = r'["""''「」]'
    
    # 计数器，用于判断奇偶
    counter = 0
    
    def replace_quote(match):
        nonlocal counter
        counter += 1
        # 奇数用中文左引号"，偶数用中文右引号"
        return '“' if counter % 2 == 1 else '”'
    # 执行替换
    result = re.sub(quote_pattern, replace_quote, content)
    
    return result, counter


def fix_punctuation(content):
    """
    将英文标点符号转换为中文标点符号
    """
    # 英文标点到中文标点的映射
    punctuation_map = {
        ',': '，',   # 逗号
        ':': '：',   # 冒号
        ';': '；',   # 分号
        '!': '！',   # 感叹号
        '?': '？',   # 问号
        '(': '（',   # 左括号
        ')': '）',   # 右括号
    }
    
    result = content
    replacement_count = 0
    
    # 逐个替换标点符号
    for eng_punct, chn_punct in punctuation_map.items():
        # 转义特殊字符
        escaped_punct = re.escape(eng_punct)
        count = len(re.findall(escaped_punct, result))
        replacement_count += count
        result = re.sub(escaped_punct, chn_punct, result)
    
    return result, replacement_count


def fix_units(content):
    """
    将中文单位转换为标准符号单位
    """
    # 中文单位到符号的映射（按长度排序，优先匹配长的）
    units_map = {
        # 面积单位
        '平方公里': 'km²',
        '平方千米': 'km²',
        '平方米': 'm²',
        '平方厘米': 'cm²',
        '平方毫米': 'mm²',
        
        # 体积单位
        '立方米': 'm³',
        '立方厘米': 'cm³',
        '立方毫米': 'mm³',
        '立方公里': 'km³',
        '立方千米': 'km³',
        
        # 长度单位
        '公里': 'km',
        '千米': 'km',
        '厘米': 'cm',
        '毫米': 'mm',
        '微米': 'μm',
        '纳米': 'nm',
        
        # 质量单位
        '公斤': 'kg',
        '千克': 'kg',
        '毫克': 'mg',
        '微克': 'μg',
        
        # 容量单位
        '毫升': 'mL',
        '微升': 'μL',
        
        # 时间单位
        '小时': 'h',
        '分钟': 'min',
        '秒钟': 's',
        
        # 温度单位
        '摄氏度': '℃',
        '华氏度': '℉',
        
        # 直接替换 m2/m3 和 km2/km3
        'km2': 'km²',
        'km3': 'km³',
        'm2': 'm²',
        'm3': 'm³',
    }
    
    result = content
    replacement_count = 0
    
    # 按长度从长到短排序，避免误匹配（如"平方米"要在"米"之前）
    sorted_units = sorted(units_map.items(), key=lambda x: len(x[0]), reverse=True)
    
    for chn_unit, symbol in sorted_units:
        count = result.count(chn_unit)
        if count > 0:
            replacement_count += count
            result = result.replace(chn_unit, symbol)
    
    return result, replacement_count


def process_paragraph(paragraph, stats):
    """
    处理段落中的文本，保持格式不变
    """
    # 处理段落文本
    if paragraph.text:
        # 获取原始文本
        original_text = paragraph.text
        
        # 应用所有转换
        fixed_text, quote_count = fix_quotes(original_text)
        fixed_text, punct_count = fix_punctuation(fixed_text)
        fixed_text, unit_count = fix_units(fixed_text)
        
        # 更新统计
        stats['quotes'] += quote_count
        stats['punctuation'] += punct_count
        stats['units'] += unit_count
        
        # 如果文本有变化，更新段落
        if fixed_text != original_text:
            # 保持格式的方式更新文本
            # 遍历所有runs，只修改第一个非空run，清空其他runs
            runs = paragraph.runs
            if runs:
                # 将所有文本放到第一个run中
                runs[0].text = fixed_text
                # 清空其他runs
                for i in range(1, len(runs)):
                    runs[i].text = ''


def process_table(table, stats):
    """
    处理表格中的文本
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph, stats)


def process_docx(input_file):
    """
    处理DOCX文件
    """
    input_path = Path(input_file)
    
    if not input_path.exists():
        print(f"❌ 错误：文件不存在 - {input_file}")
        return False
    
    if input_path.suffix.lower() != '.docx':
        print(f"❌ 错误：文件必须是.docx格式")
        return False
    
    # 生成输出文件名
    output_path = input_path.parent / f"{input_path.stem}_fixed{input_path.suffix}"
    
    try:
        # 读取文档
        print(f"📖 正在读取文件: {input_path.name}")
        doc = Document(input_path)
        
        # 统计信息
        stats = {
            'quotes': 0,
            'punctuation': 0,
            'units': 0
        }
        
        # 处理所有段落
        print(f"🔄 正在处理段落...")
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph, stats)
        
        # 处理所有表格
        print(f"🔄 正在处理表格...")
        for table in doc.tables:
            process_table(table, stats)
        
        # 处理页眉页脚
        if SKIP_FOOTER:
            # 彻底删除页眉页脚
            print(f"🗑️  正在删除页眉页脚...")
            from docx.oxml.ns import qn
            for section in doc.sections:
                sectPr = section._sectPr
                # 删除所有页眉引用
                for headerRef in sectPr.findall(qn('w:headerReference')):
                    sectPr.remove(headerRef)
                # 删除所有页脚引用
                for footerRef in sectPr.findall(qn('w:footerReference')):
                    sectPr.remove(footerRef)
        else:
            # 处理页眉页脚（格式修复）
            print(f"🔄 正在处理页眉页脚...")
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        process_paragraph(paragraph, stats)
                    for table in section.header.tables:
                        process_table(table, stats)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        process_paragraph(paragraph, stats)
                    for table in section.footer.tables:
                        process_table(table, stats)
        
        # 保存文件
        print(f"💾 正在保存文件...")
        doc.save(output_path)
        
        print(f"✅ 处理完成！")
        print(f"   - 共替换了 {stats['quotes']} 个引号")
        print(f"   - 共替换了 {stats['punctuation']} 个标点符号")
        print(f"   - 共转换了 {stats['units']} 个单位")
        print(f"   - 输出文件: {output_path.name}")
        
        return True
        
    except Exception as e:
        print(f"❌ 处理失败：{e}")
        import traceback
        traceback.print_exc()
        return False


def get_finder_selection():
    """获取 Finder 中选中的文件"""
    import subprocess
    script = '''
    tell application "Finder"
        set theSelection to selection
        if (count of theSelection) > 0 then
            return POSIX path of (item 1 of theSelection as alias)
        else
            return ""
        end if
    end tell
    '''
    try:
        result = subprocess.run(['osascript', '-e', script], 
                              capture_output=True, text=True, timeout=5)
        return result.stdout.strip()
    except:
        return ""


if __name__ == "__main__":
    input_file = None
    
    # 优先使用命令行参数
    if len(sys.argv) >= 2 and sys.argv[1].strip():
        input_file = sys.argv[1].strip()
    else:
        # 尝试从 Finder 获取选中的文件
        finder_file = get_finder_selection()
        if finder_file and finder_file.lower().endswith('.docx'):
            input_file = finder_file
            print(f"📂 从 Finder 获取文件: {Path(finder_file).name}")
    
    if not input_file:
        print("❌ 错误：缺少文件名参数")
        print("\n使用方法：")
        print("  1. 在 Finder 中选中 .docx 文件，然后运行此脚本")
        print("  2. 或在 Raycast 中输入文件路径")
        print("\n示例：")
        print("    python3 docx_text_formatter.py 文件名.docx")
        sys.exit(1)
    
    print("=" * 50)
    print("文本格式自动修复工具 - DOCX版本")
    print("=" * 50)
    
    success = process_docx(input_file)
    
    if success:
        print("\n🎉 全部完成！")
    else:
        print("\n❌ 处理失败，请检查错误信息")
        sys.exit(1)

