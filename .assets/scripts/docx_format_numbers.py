#!/Users/tianli/miniforge3/bin/python3
# @raycast.schemaVersion 1
# @raycast.title docx-numbers
# @raycast.mode fullOutput
# @raycast.icon 📄
# @raycast.packageName Scripts
# @raycast.description Format numbers font
# -*- coding: utf-8 -*-
"""
Word文档数字和英文字母字体格式化工具

功能：
- 将文档中所有数字和英文字母设置为 Times New Roman 字体
- 保持中文字符原有字体
- 保持所有格式属性（粗体、斜体、颜色等）
- 处理段落和表格中的内容
- 自动备份原文件
"""

import sys
import os
import re
import shutil
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from common_utils import get_input_files

def show_message(msg_type, message):
    """显示格式化消息"""
    icons = {
        'success': '✅',
        'error': '❌',
        'warning': '⚠️',
        'info': 'ℹ️',
        'processing': '🔄'
    }
    icon = icons.get(msg_type, 'ℹ️')
    print(f"{icon} {message}")


def backup_file(file_path):
    """备份原始文件"""
    backup_path = f"{file_path}.backup"
    try:
        shutil.copy2(file_path, backup_path)
        show_message('info', f"已备份原文件: {backup_path}")
        return backup_path
    except Exception as e:
        show_message('warning', f"备份文件失败: {e}")
        return None


def is_ascii_char(char):
    """判断字符是否为数字或英文字母"""
    return bool(re.match(r'[0-9a-zA-Z]', char))


def copy_run_format(source_run, target_run):
    """复制run的格式属性"""
    # 复制字体属性
    if source_run.bold is not None:
        target_run.bold = source_run.bold
    if source_run.italic is not None:
        target_run.italic = source_run.italic
    if source_run.underline is not None:
        target_run.underline = source_run.underline
    
    # 复制字体大小
    if source_run.font.size is not None:
        target_run.font.size = source_run.font.size
    
    # 复制字体颜色
    if source_run.font.color.rgb is not None:
        target_run.font.color.rgb = source_run.font.color.rgb
    
    # 复制其他属性
    if source_run.font.highlight_color is not None:
        target_run.font.highlight_color = source_run.font.highlight_color
    
    if source_run.font.subscript is not None:
        target_run.font.subscript = source_run.font.subscript
    
    if source_run.font.superscript is not None:
        target_run.font.superscript = source_run.font.superscript


def get_original_east_asia_font(run):
    """
    获取run原有的中文字体
    """
    rPr = run._element.rPr
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            east_asia = rFonts.get(qn('w:eastAsia'))
            if east_asia:
                return east_asia
    return None


def set_font_for_run(run, font_name='Times New Roman', is_ascii=True, preserve_east_asia=True, original_east_asia=None):
    """
    为run设置字体，最强制的方式
    """
    if not is_ascii:
        return
    
    # 1. 使用API设置
    run.font.name = font_name
    
    # 2. 强制XML级别设置
    rPr = run._element.get_or_add_rPr()
    
    # 清除可能存在的rStyle（样式引用）
    rStyle = rPr.find(qn('w:rStyle'))
    if rStyle is not None:
        rPr.remove(rStyle)
    
    # 获取或创建rFonts
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    
    # 强制清除所有字体设置，重新设置
    for attr in [qn('w:ascii'), qn('w:hAnsi'), qn('w:cs'), qn('w:eastAsia'), 
                 qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:cstheme'), qn('w:eastAsiaTheme')]:
        if attr in rFonts.attrib:
            del rFonts.attrib[attr]
    
    # 设置新字体
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    
    # 保留中文字体
    if original_east_asia:
        rFonts.set(qn('w:eastAsia'), original_east_asia)


def split_text_by_type(text):
    """
    将文本按字符类型分段
    返回: [(text_segment, is_ascii), ...]
    """
    if not text:
        return []
    
    segments = []
    current_segment = text[0]
    current_is_ascii = is_ascii_char(text[0])
    
    for char in text[1:]:
        char_is_ascii = is_ascii_char(char)
        
        if char_is_ascii == current_is_ascii:
            current_segment += char
        else:
            segments.append((current_segment, current_is_ascii))
            current_segment = char
            current_is_ascii = char_is_ascii
    
    # 添加最后一段
    segments.append((current_segment, current_is_ascii))
    
    return segments


def process_run(run, paragraph):
    """
    处理单个run，将数字和英文字母设置为Times New Roman
    """
    if not run.text:
        return
    
    # 检查文本中是否有数字或英文字母
    has_ascii = bool(re.search(r'[0-9a-zA-Z]', run.text))
    if not has_ascii:
        # 纯中文，不需要修改
        return
    
    # 获取原有的中文字体设置
    original_east_asia = get_original_east_asia_font(run)
    
    # 分析文本段落
    segments = split_text_by_type(run.text)
    
    # 如果只有一种类型（纯ASCII），直接处理
    if len(segments) == 1:
        text, is_ascii = segments[0]
        if is_ascii:
            # 纯数字或英文，设置为Times New Roman
            set_font_for_run(run, 'Times New Roman', is_ascii=True, 
                           original_east_asia=original_east_asia)
        return
    
    # 需要拆分run - 混合内容
    # 保存原run的所有格式属性
    original_text = run.text
    original_format = {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'size': run.font.size,
        'color': run.font.color.rgb if run.font.color.rgb else None,
        'highlight': run.font.highlight_color,
        'subscript': run.font.subscript,
        'superscript': run.font.superscript,
    }
    
    # 获取原run的位置
    run_element = run._element
    parent = run_element.getparent()
    run_index = list(parent).index(run_element)
    
    # 清空原run的文本（保留它作为第一个segment的容器）
    first_segment, first_is_ascii = segments[0]
    run.text = first_segment
    
    # 设置第一个segment的字体
    if first_is_ascii:
        set_font_for_run(run, 'Times New Roman', is_ascii=True,
                        original_east_asia=original_east_asia)
    else:
        # 即使是中文，也确保格式被正确复制（原run已经被修改，需要重新应用格式）
        if original_format['bold'] is not None:
            run.bold = original_format['bold']
        if original_format['italic'] is not None:
            run.italic = original_format['italic']
        if original_format['underline'] is not None:
            run.underline = original_format['underline']
        if original_format['size']:
            run.font.size = original_format['size']
        if original_format['color']:
            run.font.color.rgb = original_format['color']
        if original_format['highlight']:
            run.font.highlight_color = original_format['highlight']
        if original_format['subscript']:
            run.font.subscript = original_format['subscript']
        if original_format['superscript']:
            run.font.superscript = original_format['superscript']
    
    # 为剩余的segments创建新的runs（从第二个开始）
    current_index = run_index
    for text, is_ascii in segments[1:]:
        new_run = paragraph.add_run(text)
        
        # 复制格式
        if original_format['bold'] is not None:
            new_run.bold = original_format['bold']
        if original_format['italic'] is not None:
            new_run.italic = original_format['italic']
        if original_format['underline'] is not None:
            new_run.underline = original_format['underline']
        if original_format['size']:
            new_run.font.size = original_format['size']
        if original_format['color']:
            new_run.font.color.rgb = original_format['color']
        if original_format['highlight']:
            new_run.font.highlight_color = original_format['highlight']
        if original_format['subscript']:
            new_run.font.subscript = original_format['subscript']
        if original_format['superscript']:
            new_run.font.superscript = original_format['superscript']
        
        # 设置字体
        if is_ascii:
            set_font_for_run(new_run, 'Times New Roman', is_ascii=True,
                           original_east_asia=original_east_asia)
        
        # 将新run插入到正确位置
        current_index += 1
        parent.insert(current_index, new_run._element)


def process_paragraph(paragraph, stats=None):
    """处理段落中的所有run"""
    if stats is None:
        stats = {'processed_runs': 0, 'ascii_runs': 0}
    
    # 需要创建run列表的副本，因为处理过程中可能会修改列表
    runs = list(paragraph.runs)
    
    # 正序处理（从前往后）
    for run in runs:
        try:
            # 检查run是否仍然存在（可能在之前的处理中被删除）
            if run._element.getparent() is None:
                continue
            
            # 检查是否包含ASCII字符
            if run.text and re.search(r'[0-9a-zA-Z]', run.text):
                stats['ascii_runs'] += 1
                process_run(run, paragraph)
                stats['processed_runs'] += 1
        except Exception as e:
            show_message('warning', f"处理run时出错: {e}")
            import traceback
            traceback.print_exc()
            continue
    
    return stats


def process_table(table, stats=None):
    """处理表格中的所有单元格"""
    if stats is None:
        stats = {'processed_runs': 0, 'ascii_runs': 0}
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph, stats)
    
    return stats


def format_document(input_path, output_path=None):
    """
    格式化Word文档中的数字和英文字母
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径，如果为None则覆盖原文件
    
    Returns:
        bool: 是否成功
    """
    try:
        # 验证输入文件
        if not os.path.exists(input_path):
            show_message('error', f"文件不存在: {input_path}")
            return False
        
        if not input_path.lower().endswith('.docx'):
            show_message('error', "只支持.docx格式的文件")
            return False
        
        show_message('processing', f"正在处理文件: {os.path.basename(input_path)}")
        
        # 备份原文件
        backup_path = backup_file(input_path)
        
        # 打开文档
        doc = Document(input_path)
        
        # 统计信息
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)
        stats = {'processed_runs': 0, 'ascii_runs': 0}
        
        show_message('info', f"文档包含 {total_paragraphs} 个段落，{total_tables} 个表格")
        
        # 处理所有段落
        show_message('processing', "正在处理段落...")
        for i, paragraph in enumerate(doc.paragraphs, 1):
            try:
                process_paragraph(paragraph, stats)
            except Exception as e:
                show_message('warning', f"处理第{i}个段落时出错: {e}")
                continue
        
        # 处理所有表格
        if total_tables > 0:
            show_message('processing', "正在处理表格...")
            for i, table in enumerate(doc.tables, 1):
                try:
                    process_table(table, stats)
                except Exception as e:
                    show_message('warning', f"处理第{i}个表格时出错: {e}")
                    continue
        
        show_message('info', f"已处理 {stats['ascii_runs']} 个包含数字/英文的run")
        
        # 保存文档
        output_file = output_path if output_path else input_path
        doc.save(output_file)
        
        show_message('success', f"格式化完成: {os.path.basename(output_file)}")
        
        if backup_path:
            show_message('info', f"如需恢复，请使用备份文件: {os.path.basename(backup_path)}")
        
        return True
        
    except Exception as e:
        show_message('error', f"处理文件时出错: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """主函数"""
    # 获取输入文件（优先命令行参数，否则从 Finder 获取）
    files = get_input_files(sys.argv[1:], expected_ext='docx', allow_multiple=False)
    
    if not files:
        print("用法:")
        print(f"  python {os.path.basename(__file__)} <输入文件.docx>")
        print(f"  python {os.path.basename(__file__)} <输入文件.docx> <输出文件.docx>")
        print("  或在 Finder 中选择 .docx 文件后运行")
        print()
        print("功能:")
        print("  将Word文档中所有数字(0-9)和英文字母(a-z, A-Z)设置为Times New Roman字体")
        print("  中文字符保持原有字体不变")
        print()
        print("示例:")
        print(f"  python {os.path.basename(__file__)} document.docx")
        print(f"  python {os.path.basename(__file__)} input.docx output.docx")
        sys.exit(1)
    
    input_path = files[0]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    # 处理文件
    success = format_document(input_path, output_path)
    
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()

